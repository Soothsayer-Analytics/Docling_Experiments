import os
import logging
from typing import List, Dict, Any, Optional, Tuple, TypedDict, Union
from dataclasses import dataclass
from pathlib import Path
from datetime import datetime
import json
import time
import streamlit as st
import re
from document_chunk import DocumentChunk

# Core libraries
from openai import AzureOpenAI
import fitz  # PyMuPDF
from docx import Document
import base64

# Vector store and embeddings
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import pickle

# LangGraph imports
from langgraph.graph import StateGraph, END
from langgraph.graph.message import add_messages
from langchain_core.messages import HumanMessage, AIMessage

from typing_extensions import Annotated

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# File grouping configuration
FILE_GROUPS = {
    "Chemical Consumption": [
        "chemical", "consumption", "chemicals"
    ],
    "Engineering Tickets": [
        "engineering", "ticket", "tickets"
    ],
    "Risk Assessment and Hazard Analysis": [
        "risk", "hazard", "safety", "assessment", "analysis", "incident", "accident"
    ],
    "Well Recap": [
        "well", "recap"
    ],
    "Mud Program": [
        "mud", "program"
    ],
    "Contractor Feedback": [
        "contractor", "feedback"
    ],
    "Hydraulic Summary": [
        "hydraulic", "summary","hydraulics"
    ],
    "Other Group": []  # Default group for unmatched files
}

@dataclass
class FileSystemInfo:
    """Stores file system structure information"""
    total_folders: int
    total_files: int
    file_types: Dict[str, int]
    folder_structure: Dict[str, Any]
    file_groups: Dict[str, List[str]]

# LangGraph State
class AgentState(TypedDict):
    messages: Annotated[List[Union[HumanMessage, AIMessage]], add_messages]
    question: str
    selected_groups: List[str]
    search_results: List[Dict[str, Any]]
    final_answer: str
    confidence: float

class DocumentProcessor:
    """Enhanced document processor with method tracking and validation"""
    
    def __init__(self, azure_client: AzureOpenAI):
        """Initialize with Azure client and setup processing stats"""
        if not azure_client:
            raise ValueError("AzureOpenAI client cannot be None")
            
        self.azure_client = azure_client
        self.supported_extensions = {'.pdf', '.docx', '.txt'}
        self.processing_stats = {
            'text_extraction': [],
            'vision_ocr': [],
            'total_processed': 0
        }
        self.metadata_summary = []  # Stores document metadata for quick lookup
        self.summary_file = "document_summary.xlsx"
    
    def classify_file_group(self, file_path: str) -> str:
        """Classify file into appropriate group based on filename"""
        file_name = os.path.basename(file_path).lower()
        
        for group_name, keywords in FILE_GROUPS.items():
            if group_name == "Other Group":
                continue
            
            for keyword in keywords:
                if keyword.lower() in file_name:
                    return group_name
        
        return "Other Group"
    
    def process_pdf(self, file_path: str, ocr_threshold: int = 50) -> List[DocumentChunk]:
        """Process PDF file, extracting text and using OCR when needed"""
        if not os.path.exists(file_path) or os.path.getsize(file_path) == 0:
            logger.warning(f"Invalid PDF file: {file_path}")
            return []

        full_text = ""
        vision_pages = []
        text_pages = []
        file_group = self.classify_file_group(file_path)

        try:
            with fitz.open(file_path) as doc:
                total_pages = len(doc)
                for page_num in range(min(total_pages, 50)):
                    page = doc[page_num]
                    page_text = page.get_text()
                    
                    if page_text and len(page_text.strip()) >= ocr_threshold:
                        full_text += page_text + "\n"
                        text_pages.append(page_num + 1)
                    else:
                        ocr_text = self._extract_text_via_ocr_page(page)
                        full_text += ocr_text + "\n"
                        vision_pages.append(page_num + 1)

            chunks = self._create_chunks(full_text, file_path, "pdf", file_group)
            if not chunks:
                return []

            # Track processing stats
            primary_method = 'Vision (OCR)' if vision_pages else 'Text Extraction'
            self.processing_stats['vision_ocr' if vision_pages else 'text_extraction'].append({
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'group': file_group,
                'total_pages': total_pages,
                'text_pages': text_pages,
                'vision_pages': vision_pages,
                'primary_method': primary_method
            })
            self.processing_stats['total_processed'] += 1

            chunks[0].metadata.update({
                "processing_method": primary_method,
                "text_pages": text_pages,
                "vision_pages": vision_pages,
                "total_pages": total_pages
            })
            return chunks

        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return []


    def get_processing_stats(self) -> Dict[str, Any]:
        """Get detailed processing statistics"""
        return {
            'total_processed': self.processing_stats['total_processed'],
            'text_extraction_count': len(self.processing_stats['text_extraction']),
            'vision_ocr_count': len(self.processing_stats['vision_ocr']),
            'text_extraction_files': self.processing_stats['text_extraction'],
            'vision_ocr_files': self.processing_stats['vision_ocr'],
            'document_summaries': len(self.metadata_summary)
        }

    def _generate_summary(self, text: str, file_name: str) -> str:
        """Generate concise summary of document content"""
        try:
            prompt = f"""Create a concise 1-2 sentence summary of this document focusing on key wells, dates, and operations.
            Be specific about any numerical data, well names, or important dates.
            Document: {file_name}
            Content: {text[:5000]}... [truncated if long]
            
            Summary:"""
            
            response = self.azure_client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=200,
                temperature=0.0
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Error generating summary for {file_name}: {e}")
            return "Summary generation failed"

    def _save_summary_to_excel(self):
        """Save document summaries to Excel file"""
        try:
            import pandas as pd
            df = pd.DataFrame(self.metadata_summary)
            
            # Reorder columns for better readability
            df = df[['file_name', 'group', 'file_type', 'dates', 'wells', 'summary', 'processed_at']]
            
            # Save to Excel
            df.to_excel(self.summary_file, index=False)
            logger.info(f"Updated document summary Excel file: {self.summary_file}")
        except Exception as e:
            logger.error(f"Error saving summary Excel: {e}")
    
    def _extract_text_via_ocr_page(self, page) -> str:
        """Extract text from a single PDF page using GPT-4V OCR"""
        try:
            if len(page.get_text().strip()) > 50:
                return page.get_text()
            
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            img_data = pix.tobytes("png")
            img_base64 = base64.b64encode(img_data).decode()
            
            response = self.azure_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Extract all text from this image. Maintain the original formatting as much as possible. If there are tables, preserve their structure. Return only the extracted text without any commentary."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_base64}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=4000,
                temperature=0
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            logger.error(f"OCR failed for page: {e}")
            return ""
    
    def process_docx(self, file_path: str) -> List[DocumentChunk]:
        """Process DOCX files into chunks"""
        try:
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            
            for table in doc.tables:
                text += "\n" + "\n".join(
                    " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    for row in table.rows
                )
            
            return self._create_chunks(text, file_path, "docx", self.classify_file_group(file_path))
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return []
    
    def process_txt(self, file_path: str) -> List[DocumentChunk]:
        """Process TXT files into chunks"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return self._create_chunks(file.read(), file_path, "txt", self.classify_file_group(file_path))
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            return []
    
    def _extract_metadata(self, text: str) -> dict:
        """Extract key metadata like dates and well names from text"""
        # Date patterns (dd/mm/yyyy, mm/dd/yyyy, etc.)
        date_pattern = r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"
        dates = list(set(re.findall(date_pattern, text)))
        
        # Well name patterns (Well XYZ, Well-123, etc.)
        well_pattern = r"\bWell\s*[A-Z0-9-]+\b"
        wells = list(set(re.findall(well_pattern, text, re.IGNORECASE)))
        
        return {
            "dates": sorted(dates),
            "wells": wells
        }

    def _create_chunks(self, text: str, file_path: str, file_type: str, group: str) -> List[DocumentChunk]:
        """Split text into chunks with metadata and generate document summary"""
        if not text.strip():
            return []

        chunks = []
        current_chunk = ""
        chunk_size = 800
        overlap = 100
        file_name = os.path.basename(file_path)
        
        # Extract metadata for summary
        metadata = self._extract_metadata(text)
        
        # Generate document summary using GPT
        summary = self._generate_summary(text, file_name)
        
        # Add to enhanced metadata summary
        self.metadata_summary.append({
            "file_name": file_name,
            "group": group,
            "dates": metadata["dates"],
            "wells": metadata["wells"],
            "summary": summary,
            "processed_at": datetime.now().isoformat(),
            "file_type": file_type
        })
        
        # Save updated summary to Excel
        self._save_summary_to_excel()
        
        def create_chunk(content: str, idx: int) -> DocumentChunk:
            return DocumentChunk(
                content=content.strip(),
                source_file=file_path,
                chunk_id=f"{file_name}_chunk_{idx}",
                group=group,
                metadata={
                    "file_type": file_type,
                    "chunk_index": idx,
                    "file_size": len(content),
                    "processed_at": datetime.now().isoformat(),
                    "file_name": file_name,
                    "group": group,
                    "processing_method": "Text Extraction" if file_type != "pdf" else "",
                    "dates": metadata["dates"],
                    "wells": metadata["wells"]
                }
            )

        sentences = [s.strip() for s in text.split('. ') if s.strip()]
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) + 2 > chunk_size and current_chunk:
                chunks.append(create_chunk(current_chunk, len(chunks)))
                # Keep overlap words
                words = current_chunk.split()
                current_chunk = " ".join(words[-overlap//10:]) + " " + sentence + ". "
            else:
                current_chunk += sentence + ". "

        if current_chunk.strip():
            chunks.append(create_chunk(current_chunk, len(chunks)))

        logger.info(f"Created {len(chunks)} chunks from {file_path} in group {group}")
        return chunks

class GroupedVectorStore:
    """Manages document embeddings with group-based organization"""
    
    def __init__(self, embedding_model_name: str = "all-MiniLM-L6-v2"):
        self.embedding_model = SentenceTransformer(embedding_model_name)
        self.group_indices = {}
        self.group_chunks = {}
        self.dimension = 384
        
        # Enhanced tracking
        self.file_type_stats = {}
        self.group_file_details = {}
        self.date_index = {}  # {date: [file1, file2]}
        self.well_index = {}  # {well_name: [file1, file2]}
        
        logger.info(f"Initialized GroupedVectorStore with model: {embedding_model_name}")
        
    def _update_indices(self, chunk: DocumentChunk):
        """Update date and well indices"""
        file_name = chunk.metadata.get('file_name', 'unknown')
        dates = chunk.metadata.get('dates', [])
        wells = chunk.metadata.get('wells', [])
        
        for date in dates:
            if date not in self.date_index:
                self.date_index[date] = []
            if file_name not in self.date_index[date]:
                self.date_index[date].append(file_name)
                
        for well in wells:
            if well not in self.well_index:
                self.well_index[well] = []
            if file_name not in self.well_index[well]:
                self.well_index[well].append(file_name)
    
    def add_documents(self, chunks: List[DocumentChunk]):
        """Enhanced document addition with detailed tracking"""
        if not chunks:
            logger.warning("No chunks provided to add_documents")
            return
        
        # Track file types and group details
        for chunk in chunks:
            file_type = chunk.metadata.get('file_type', 'unknown')
            group = chunk.group
            file_name = chunk.metadata.get('file_name', 'unknown')
            
            # Update file type stats
            if file_type not in self.file_type_stats:
                self.file_type_stats[file_type] = 0
            self.file_type_stats[file_type] += 1
            
            # Update group file details
            if group not in self.group_file_details:
                self.group_file_details[group] = {
                    'files': set(),
                    'file_details': {},
                    'total_chunks': 0
                }
            
            self.group_file_details[group]['files'].add(file_name)
            self.group_file_details[group]['total_chunks'] += 1
            
            if file_name not in self.group_file_details[group]['file_details']:
                self.group_file_details[group]['file_details'][file_name] = {
                    'file_type': file_type,
                    'processing_method': chunk.metadata.get('processing_method', 'Unknown'),
                    'chunks': 0,
                    'total_pages': chunk.metadata.get('total_pages', 0),
                    'text_pages': chunk.metadata.get('text_pages', []),
                    'vision_pages': chunk.metadata.get('vision_pages', []),
                    'dates': chunk.metadata.get('dates', []),
                    'wells': chunk.metadata.get('wells', [])
                }
            
            self.group_file_details[group]['file_details'][file_name]['chunks'] += 1
            
            # Update date and well indices
            self._update_indices(chunk)
        
        # Group chunks by their group classification
        grouped_chunks = {}
        for chunk in chunks:
            group = chunk.group
            if group not in grouped_chunks:
                grouped_chunks[group] = []
            grouped_chunks[group].append(chunk)
        
        # Process each group
        for group_name, group_chunks in grouped_chunks.items():
            logger.info(f"Adding {len(group_chunks)} chunks to group: {group_name}")
            self._add_to_group(group_name, group_chunks)
    
    def _add_to_group(self, group_name: str, chunks: List[DocumentChunk]):
        """Add chunks to a group's vector store with batch embeddings"""
        if not chunks:
            return

        # Initialize index if needed
        if group_name not in self.group_indices:
            self.group_indices[group_name] = faiss.IndexFlatIP(self.dimension)
            self.group_chunks[group_name] = []

        # Process in batches
        batch_size = 32
        embeddings = []
        
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            try:
                batch_embeddings = self.embedding_model.encode(
                    [chunk.content for chunk in batch],
                    show_progress_bar=False
                )
                embeddings.extend(batch_embeddings)
                
                # Store embeddings in chunks
                for chunk, emb in zip(batch, batch_embeddings):
                    chunk.embedding = emb
                    
            except Exception as e:
                logger.error(f"Embedding generation failed for batch {i//batch_size}: {e}")
                continue

        if embeddings:
            # Add to FAISS index
            self.group_indices[group_name].add(np.array(embeddings).astype('float32'))
            self.group_chunks[group_name].extend(chunks)
    
    def search_groups(self, query: str, groups: List[str], k: int = 5) -> List[Tuple[DocumentChunk, float, str]]:
        """Search across specified groups"""
        all_results = []
        
        for group_name in groups:
            if group_name not in self.group_indices or len(self.group_chunks[group_name]) == 0:
                logger.warning(f"Group {group_name} is empty or not initialized")
                continue
            
            try:
                query_embedding = self.embedding_model.encode([query])
                scores, indices = self.group_indices[group_name].search(
                    query_embedding.astype('float32'), 
                    min(k, len(self.group_chunks[group_name]))
                )
                
                for score, idx in zip(scores[0], indices[0]):
                    if 0 <= idx < len(self.group_chunks[group_name]):
                        all_results.append((
                            self.group_chunks[group_name][idx], 
                            float(score), 
                            group_name
                        ))
                        
            except Exception as e:
                logger.error(f"Error searching group {group_name}: {e}")
                continue
        
        # Sort by score and return top k
        all_results.sort(key=lambda x: x[1], reverse=True)
        logger.info(f"Found {len(all_results)} results across groups {groups}")
        return all_results[:k]
    
    def get_enhanced_group_stats(self) -> Dict[str, Any]:
        """Get enhanced statistics with file details"""
        stats = {
            'file_type_distribution': dict(self.file_type_stats),
            'total_groups': len(self.group_indices),
            'groups': {}
        }
        
        for group_name in self.group_indices:
            group_details = self.group_file_details.get(group_name, {})
            chunk_count = len(self.group_chunks.get(group_name, []))
            
            stats['groups'][group_name] = {
                'chunk_count': chunk_count,
                'unique_files': len(group_details.get('files', set())),
                'file_list': list(group_details.get('files', set())),
                'file_details': group_details.get('file_details', {}),
                'total_chunks_tracked': group_details.get('total_chunks', 0)
            }
        
        return stats
    
    def get_file_type_summary(self) -> Dict[str, Any]:
        """Get summary by file type"""
        summary = {
            'pdf_files': 0,
            'docx_files': 0,
            'txt_files': 0,
            'total_files': 0,
            'processing_methods': {
                'text_extraction': 0,
                'vision_ocr': 0
            }
        }
        
        for group_name, group_details in self.group_file_details.items():
            for file_name, file_info in group_details.get('file_details', {}).items():
                file_type = file_info.get('file_type', 'unknown')
                processing_method = file_info.get('processing_method', 'Unknown')
                
                if file_type == 'pdf':
                    summary['pdf_files'] += 1
                elif file_type == 'docx':
                    summary['docx_files'] += 1
                elif file_type == 'txt':
                    summary['txt_files'] += 1
                
                summary['total_files'] += 1
                
                if 'Text Extraction' in processing_method:
                    summary['processing_methods']['text_extraction'] += 1
                elif 'Vision (OCR)' in processing_method:
                    summary['processing_methods']['vision_ocr'] += 1
        
        return summary
    

    def get_group_files(self, group_name: str) -> Dict[str, Any]:
        """Get detailed information about files in a specific group"""
        if group_name not in self.group_file_details:
            return {'files': [], 'total_files': 0}
        
        group_info = self.group_file_details[group_name]
        return {
            'files': list(group_info.get('files', set())),
            'total_files': len(group_info.get('files', set())),
            'file_details': group_info.get('file_details', {}),
            'total_chunks': group_info.get('total_chunks', 0)
        }
    
    def save_group(self, group_name: str, folder_path: str):
        """Save a specific group's vector store to a folder"""
        try:
            os.makedirs(folder_path, exist_ok=True)
            
            if group_name in self.group_indices:
                index_path = os.path.join(folder_path, f"{group_name}.index")
                faiss.write_index(self.group_indices[group_name], index_path)
                logger.info(f"Saved FAISS index for group {group_name} to {index_path}")
            
            if group_name in self.group_chunks:
                chunks_to_save = []
                for chunk in self.group_chunks[group_name]:
                    chunk_copy = DocumentChunk(
                        content=chunk.content,
                        source_file=chunk.source_file,
                        chunk_id=chunk.chunk_id,
                        group=chunk.group,
                        metadata=chunk.metadata,
                        embedding=None
                    )
                    chunks_to_save.append(chunk_copy)
                
                chunks_path = os.path.join(folder_path, f"{group_name}.chunks")
                with open(chunks_path, 'wb') as f:
                    pickle.dump(chunks_to_save, f)
                
                logger.info(f"Saved {len(chunks_to_save)} chunks for group {group_name} to {chunks_path}")
                
        except Exception as e:
            logger.error(f"Error saving group {group_name}: {e}")
    
    def save_all_groups(self, folder_path: str):
        """Save all groups to a single folder"""
        try:
            os.makedirs(folder_path, exist_ok=True)
            for group_name in self.group_indices:
                self.save_group(group_name, folder_path)
            logger.info(f"Saved all groups to folder: {folder_path}")
            return True
        except Exception as e:
            logger.error(f"Error saving all groups: {e}")
            return False
    
    def load_group(self, group_name: str, folder_path: str):
        """Load a specific group's vector store from folder"""
        try:
            index_path = os.path.join(folder_path, f"{group_name}.index")
            chunks_path = os.path.join(folder_path, f"{group_name}.chunks")
            
            if os.path.exists(index_path):
                self.group_indices[group_name] = faiss.read_index(index_path)
                logger.info(f"Loaded FAISS index for group {group_name} from {index_path}")
            else:
                logger.warning(f"FAISS index file not found at {index_path}")
                return False
            
            if os.path.exists(chunks_path):
                with open(chunks_path, 'rb') as f:
                    self.group_chunks[group_name] = pickle.load(f)
                logger.info(f"Loaded {len(self.group_chunks[group_name])} chunks for group {group_name} from {chunks_path}")
            else:
                logger.warning(f"Chunks file not found at {chunks_path}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error loading group {group_name}: {e}")
            return False
    
    def load_all_groups(self, folder_path: str):
        """Load all available groups from folder"""
        loaded_groups = []
        
        if not os.path.exists(folder_path):
            logger.error(f"Vector store folder not found: {folder_path}")
            return loaded_groups
            
        for filename in os.listdir(folder_path):
            if filename.endswith('.index'):
                group_name = filename[:-6]  # Remove '.index'
                if self.load_group(group_name, folder_path):
                    loaded_groups.append(group_name)
        
        logger.info(f"Loaded groups from {folder_path}: {loaded_groups}")
        return loaded_groups

class FileSystemAnalyzer:
    """Analyzes file system structure with group classification"""
    
    def __init__(self, root_path: str):
        self.root_path = Path(root_path)
        self.file_info = None
        self.document_processor = None
    
    def set_document_processor(self, processor: DocumentProcessor):
        """Set document processor for file classification"""
        self.document_processor = processor
    
    def analyze(self) -> FileSystemInfo:
        """Analyze file system structure with group classification"""
        folder_count = 0
        file_count = 0
        file_types = {}
        folder_structure = {}
        file_groups = {group: [] for group in FILE_GROUPS.keys()}
        
        try:
            for root, dirs, files in os.walk(self.root_path):
                folder_count += len(dirs)
                file_count += len(files)
                
                for file in files:
                    file_path = os.path.join(root, file)
                    ext = Path(file).suffix.lower()
                    file_types[ext] = file_types.get(ext, 0) + 1
                    
                    # Classify file into group
                    if self.document_processor:
                        group = self.document_processor.classify_file_group(file_path)
                        file_groups[group].append(file)
                
                rel_path = os.path.relpath(root, self.root_path)
                if rel_path == '.':
                    rel_path = 'root'
                
                folder_structure[rel_path] = {
                    'subdirs': dirs,
                    'files': files,
                    'file_count': len(files)
                }
        
        except Exception as e:
            logger.error(f"Error analyzing file system: {e}")
        
        self.file_info = FileSystemInfo(
            total_folders=folder_count,
            total_files=file_count,
            file_types=file_types,
            folder_structure=folder_structure,
            file_groups=file_groups
        )
        
        logger.info(f"Analyzed file system: {file_count} files, {folder_count} folders")
        return self.file_info
    
    def get_folder_summary(self) -> str:
        """Get a summary of the folder structure with group information"""
        if not self.file_info:
            self.analyze()
        
        summary = f"File System Analysis:\n"
        summary += f"- Total folders: {self.file_info.total_folders}\n"
        summary += f"- Total files: {self.file_info.total_files}\n"
        summary += f"- File types: {dict(self.file_info.file_types)}\n\n"
        
        summary += "File Groups:\n"
        for group, files in self.file_info.file_groups.items():
            summary += f"- {group}: {len(files)} files\n"
        
        return summary

class LangGraphAgent:
    """LangGraph-based agent for intelligent query processing"""
    
    def __init__(self, azure_client: AzureOpenAI, vector_store: GroupedVectorStore, file_analyzer: FileSystemAnalyzer):
        self.azure_client = azure_client
        self.vector_store = vector_store
        self.file_analyzer = file_analyzer
        self.graph = self._create_graph()
        
    def _is_date_query(self, question: str) -> bool:
        """Check if query contains date ranges and if we have matching docs"""
        date_pattern = r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"
        dates = re.findall(date_pattern, question)
        if not dates:
            return False
            
        # Check summary Excel first for matching dates
        try:
            import pandas as pd
            df = pd.read_excel("document_summary.xlsx")
            for date in dates:
                if df['dates'].astype(str).str.contains(date).any():
                    return True
        except Exception as e:
            logger.warning(f"Error checking date summaries: {e}")
            
        return bool(dates)
        
    def _is_well_query(self, question: str) -> bool:
        """Check if query contains well names and if we have matching docs"""
        well_pattern = r"\bWell\s*[A-Z0-9-]+\b"
        wells = re.findall(well_pattern, question, re.IGNORECASE)
        if not wells:
            return False
            
        # Check summary Excel first for matching wells
        try:
            import pandas as pd
            df = pd.read_excel("document_summary.xlsx")
            for well in wells:
                if df['wells'].astype(str).str.contains(well, case=False).any():
                    return True
        except Exception as e:
            logger.warning(f"Error checking well summaries: {e}")
            
        return bool(wells)

    def _get_docs_for_dates(self, dates: List[str]) -> List[str]:
        """Get documents that contain any of the specified dates"""
        try:
            import pandas as pd
            df = pd.read_excel("document_summary.xlsx")
            matching_docs = []
            for date in dates:
                matches = df[df['dates'].astype(str).str.contains(date)]['file_name'].tolist()
                matching_docs.extend(matches)
            return list(set(matching_docs))
        except Exception as e:
            logger.error(f"Error getting docs for dates: {e}")
            return []

    def _get_docs_for_wells(self, wells: List[str]) -> List[str]:
        """Get documents that contain any of the specified wells"""
        try:
            import pandas as pd
            df = pd.read_excel("document_summary.xlsx")
            matching_docs = []
            for well in wells:
                matches = df[df['wells'].astype(str).str.contains(well, case=False)]['file_name'].tolist()
                matching_docs.extend(matches)
            return list(set(matching_docs))
        except Exception as e:
            logger.error(f"Error getting docs for wells: {e}")
            return []

    def _is_metadata_query(self, question: str) -> bool:
        """Check if query is asking for metadata/statistics"""
        metadata_keywords = [
            'how many', 'total', 'count', 'number of', 'list all',
            'pdf files', 'docx files', 'txt files',
            'well recap files', 'chemical files', 'engineering files',
            'files in', 'processed using', 'text extraction', 'vision', 'ocr'
        ]
        
        question_lower = question.lower()
        return any(keyword in question_lower for keyword in metadata_keywords)
    
    def _handle_metadata_query(self, question: str) -> Dict[str, Any]:
        """Handle metadata queries with streamlined response generation"""
        question_lower = question.lower()
        stats = self.vector_store.get_enhanced_group_stats()
        file_summary = self.vector_store.get_file_type_summary()
        
        # Define response templates for different query types
        responses = {
            'pdf': lambda: f"Total PDF files: {file_summary['pdf_files']}",
            'docx': lambda: f"Total DOCX files: {file_summary['docx_files']}",
            'txt': lambda: f"Total TXT files: {file_summary['txt_files']}",
            'well recap': lambda: self._format_group_response('Well Recap', stats),
            'chemical': lambda: self._format_group_response('Chemical Consumption', stats),
            'engineering': lambda: self._format_group_response('Engineering Tickets', stats),
            'text extraction|vision|ocr': lambda: (
                f"Files processed using Text Extraction: {file_summary['processing_methods']['text_extraction']}\n"
                f"Files processed using Vision (OCR): {file_summary['processing_methods']['vision_ocr']}"
            ),
            'total files': lambda: (
                f"Total files processed: {file_summary['total_files']}\n"
                f"PDF files: {file_summary['pdf_files']}\n"
                f"DOCX files: {file_summary['docx_files']}\n"
                f"TXT files: {file_summary['txt_files']}"
            ),
            'list all': lambda: self._format_all_groups(stats)
        }

        # Find matching response template
        answer = None
        for pattern, response_fn in responses.items():
            if any(keyword in question_lower for keyword in pattern.split('|')):
                answer = response_fn()
                break

        # Default response if no specific pattern matched
        if not answer:
            answer = (
                f"System Statistics:\n"
                f"Total files: {file_summary['total_files']}\n"
                f"Total groups: {stats['total_groups']}\n"
                f"File types: PDF({file_summary['pdf_files']}), DOCX({file_summary['docx_files']}), TXT({file_summary['txt_files']})\n"
                f"Processing: Text Extraction({file_summary['processing_methods']['text_extraction']}), "
                f"Vision OCR({file_summary['processing_methods']['vision_ocr']})"
            )

        return {
            "answer": answer,
            "confidence": 1.0,
            "selected_groups": ["System Metadata"],
            "search_results": [],
            "metadata_query": True
        }

    def _format_group_response(self, group_name: str, stats: dict) -> str:
        """Format response for group-specific queries"""
        group_stats = stats['groups'].get(group_name, {})
        response = f"Total {group_name} files: {group_stats.get('unique_files', 0)}"
        if group_stats.get('file_list'):
            response += f"\nFiles: {', '.join(group_stats['file_list'])}"
        return response

    def _format_all_groups(self, stats: dict) -> str:
        """Format response for 'list all' queries"""
        response = "File distribution by groups:\n"
        for group_name, group_info in stats['groups'].items():
            if group_info['unique_files'] > 0:
                response += f"\n{group_name}: {group_info['unique_files']} files"
                response += f" ({group_info['chunk_count']} chunks)"
        return response
    
    def _create_graph(self) -> StateGraph:
        """Create LangGraph workflow"""
        workflow = StateGraph(AgentState)
        
        # Add nodes
        workflow.add_node("group_selector", self._select_groups)
        workflow.add_node("document_searcher", self._search_documents)
        workflow.add_node("answer_synthesizer", self._synthesize_answer)
        
        # Add edges
        workflow.add_edge("group_selector", "document_searcher")
        workflow.add_edge("document_searcher", "answer_synthesizer")
        workflow.add_edge("answer_synthesizer", END)
        
        # Set entry point
        workflow.set_entry_point("group_selector")
        
        return workflow.compile()

    def visualize_workflow(self):
        """Print ASCII representation of the LangGraph workflow"""
        print("\nLangGraph Workflow:")
        print("┌────────────────┐")
        print("│    Start       │")
        print("└──────┬─────────┘")
        print("       │")
        print("┌──────▼─────────┐")
        print("│ Group Selector │")
        print("└──────┬─────────┘")
        print("       │")
        print("┌──────▼─────────┐")
        print("│ Doc Searcher   │")
        print("└──────┬─────────┘")
        print("       │")
        print("┌──────▼─────────┐")
        print("│ Answer Synth   │")
        print("└──────┬─────────┘")
        print("       │")
        print("┌──────▼─────────┐")
        print("│     End        │")
        print("└────────────────┘")
        print("\nFlow: Start → Group Selector → Document Searcher → Answer Synthesizer → End")
    
    def _select_groups(self, state: AgentState) -> AgentState:
        """Select relevant groups by first checking summary file metadata"""
        question = state["question"]
        
        try:
            import pandas as pd
            df = pd.read_excel("document_summary.xlsx")
            
            # Check for matching documents in summary
            matching_groups = set()
            
            # Search summaries for question keywords
            for _, row in df.iterrows():
                # Check if question contains any well names from this doc
                doc_wells = row['wells'] if isinstance(row['wells'], list) else eval(row['wells'])
                for well in doc_wells:
                    if well.lower() in question.lower():
                        matching_groups.add(row['group'])
                
                # Check if question contains any dates from this doc  
                doc_dates = row['dates'] if isinstance(row['dates'], list) else eval(row['dates'])
                for date in doc_dates:
                    if date in question:
                        matching_groups.add(row['group'])
                
                # Check summary content
                if row['summary'] and row['summary'].lower() in question.lower():
                    matching_groups.add(row['group'])
            
            if matching_groups:
                state["selected_groups"] = list(matching_groups)
                return state
                
        except Exception as e:
            logger.warning(f"Error checking summary file: {e}")
        
        # Fallback to AI group selection if no matches in summary
        group_selection_prompt = f"""You are an expert document group selector for oil and gas operations. 
        Strictly select ONLY the most relevant document groups based on the user's question.

        Available Groups and Their Content:
        - Chemical Consumption: ONLY chemicals, additives, fluid treatment data
        - Engineering Tickets: ONLY technical equipment, mechanical issues
        - Risk Assessment: ONLY safety documents, risk analysis
        - Well Recap: ONLY well completion, drilling reports
        - Mud Program: ONLY drilling mud, fluid properties
        - Hydraulic Summary: ONLY pressure, flow, pump data

        Question: "{question}"

        Rules:
        1. Select ONLY 1 group unless question clearly spans multiple domains
        2. Be extremely precise - wrong groups waste time
        3. If unsure between groups, prefer Chemical Consumption for chemical questions
        4. Return ONLY a JSON array of group names

        Selected Groups:"""

        try:
            response = self.azure_client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": group_selection_prompt}],
                max_tokens=200,
                temperature=0
            )
            
            groups_text = response.choices[0].message.content.strip()
            selected_groups = json.loads(groups_text)
            
            # Validate selected groups
            valid_groups = [g for g in selected_groups if g in FILE_GROUPS]
            if not valid_groups:
                valid_groups = list(FILE_GROUPS.keys())  # Search all if none valid
            
            logger.info(f"Selected groups for question '{question}': {valid_groups}")
            
            state["selected_groups"] = valid_groups
            state["messages"].append(AIMessage(content=f"Selected groups: {valid_groups}"))
            
            return state
            
        except Exception as e:
            logger.error(f"Error selecting groups: {e}")
            state["selected_groups"] = list(FILE_GROUPS.keys())
            return state
    
    def _search_documents(self, state: AgentState) -> AgentState:
        """Search documents with progressive group expansion and date/well filtering"""
        question = state["question"]
        selected_groups = state["selected_groups"]
        
        # Extract dates and wells from question for filtering
        date_pattern = r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"
        dates_in_question = re.findall(date_pattern, question)
        well_pattern = r"\bWell\s*[A-Z0-9-]+\b"
        wells_in_question = re.findall(well_pattern, question, re.IGNORECASE)
        
        # First try searching only in the primary group
        primary_results = self.vector_store.search_groups(question, selected_groups, k=5)
        
        # Filter results by dates/wells if specified in question
        if dates_in_question or wells_in_question:
            filtered_results = []
            for chunk, score, group in primary_results:
                chunk_dates = chunk.metadata.get('dates', [])
                chunk_wells = chunk.metadata.get('wells', [])
                
                # Check if chunk matches date/well filters
                date_match = not dates_in_question or any(date in chunk_dates for date in dates_in_question)
                well_match = not wells_in_question or any(well in chunk_wells for well in wells_in_question)
                
                if date_match and well_match:
                    filtered_results.append((chunk, score, group))
            
            primary_results = filtered_results
        
        # If low confidence in primary group, expand to all groups
        if (not primary_results or max(r[1] for r in primary_results) < 0.7):
            all_groups = list(FILE_GROUPS.keys())
            secondary_groups = [g for g in all_groups if g not in selected_groups]
            secondary_results = self.vector_store.search_groups(question, secondary_groups, k=3)
            primary_results.extend(secondary_results)
            
            # If still low confidence, search all documents
            if (not secondary_results or max(r[1] for r in secondary_results) < 0.5):
                all_results = self.vector_store.search_groups(question, all_groups, k=10)
                primary_results.extend(all_results)
        
        # Format results with metadata
        formatted_results = []
        for chunk, score, group in primary_results:
            formatted_results.append({
                "content": chunk.content,
                "file_name": chunk.metadata.get('file_name', 'unknown'),
                "group": group,
                "score": score,
                "chunk_id": chunk.chunk_id,
                "processing_method": chunk.metadata.get('processing_method', 'Unknown'),
                "pages": chunk.metadata.get('total_pages', 0)
            })
        
        # Sort by score and keep top results
        formatted_results.sort(key=lambda x: x["score"], reverse=True)
        state["search_results"] = formatted_results[:10]
        logger.info(f"Found {len(formatted_results)} search results with progressive expansion")
        
        return state
    
    def _synthesize_answer(self, state: AgentState) -> AgentState:
        """Synthesize answer with confidence scoring and source verification"""
        question = state["question"]
        search_results = state["search_results"]
        
        if not search_results:
            state["final_answer"] = "No relevant information found in documents."
            state["confidence"] = 0.0
            return state
        
        # Prepare detailed context with metadata
        context = []
        for result in search_results:
            context.append({
                "content": result['content'],
                "file_name": result['file_name'],
                "group": result['group'],
                "score": result['score'],
                "processing_method": result.get('processing_method', 'Unknown'),
                "pages": result.get('pages', 0)
            })
        
        # Enhanced prompt with verification steps
        synthesis_prompt = f"""You are a rigorous factual assistant for oil and gas operations.
        Carefully analyze these documents to answer the question. Follow these steps:

        1. First, verify if the documents contain direct answers to the question
        2. Cross-check facts across multiple documents
        3. Only include information that appears in at least 2 documents
        4. For chemical/technical data, prioritize documents from the Chemical Consumption group
        5. If unsure, state "Information not confirmed in documents"

        Question: {question}

        Documents:
        {json.dumps(context, indent=2)}

        Rules:
        1. Be extremely precise with chemical/technical data
        2. Include exact values when available (e.g., concentrations, quantities)
        3. Always cite specific source documents
        4. If information conflicts, note the variation
        5. Format:
           - [Exact fact from documents]. 
           - Sources: [file1, file2]
           - Confidence: [High/Medium/Low]

        Answer:"""

        try:
            response = self.azure_client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": synthesis_prompt}],
                max_tokens=1000,
                temperature=0.0
            )
            
            answer = response.choices[0].message.content
            
            # Calculate confidence based on search scores
            avg_score = sum(r["score"] for r in search_results) / len(search_results)
            confidence = min(avg_score, 1.0)
            
            state["final_answer"] = answer
            state["confidence"] = confidence
            
            logger.info(f"Generated answer with confidence: {confidence}")
            
            return state
            
        except Exception as e:
            logger.error(f"Error synthesizing answer: {e}")
            state["final_answer"] = f"I encountered an error while processing your question: {str(e)}"
            state["confidence"] = 0.0
            return state
    
    def process_query(self, question: str) -> Dict[str, Any]:
        """Enhanced query processing with metadata handling and detailed workflow explanation
        
        Query Processing Workflow:
        1. Question Analysis:
           - Check if question is about system metadata (e.g., "how many PDF files?")
           - Check for date/well references using document summaries
           - Determine if specialized group selection needed

        2. Group Selection:
           - For date/well queries: Prefer Well Recap group
           - For others: Use AI to select most relevant document groups
           - Fallback to all groups if uncertain

        3. Document Search:
           - First search selected groups
           - If low confidence, expand to other groups
           - Filter results by dates/wells if specified

        4. Answer Synthesis:
           - Verify facts across multiple documents
           - Only include info confirmed in 2+ sources
           - Cite specific source documents
           - Calculate confidence score

        5. Return structured response with:
           - Final answer
           - Confidence score (0-1)
           - Groups searched
           - Relevant document excerpts
        """
        # Check if it's a metadata query
        if self._is_metadata_query(question):
            return self._handle_metadata_query(question)
        
        # Regular document search query
        initial_state = {
            "messages": [HumanMessage(content=question)],
            "question": question,
            "selected_groups": [],
            "search_results": [],
            "final_answer": "",
            "confidence": 0.0
        }
        
        try:
            result = self.graph.invoke(initial_state)
            
            return {
                "answer": result["final_answer"],
                "confidence": result["confidence"],
                "selected_groups": result["selected_groups"],
                "search_results": result["search_results"],
                "metadata_query": False,
                "workflow": "See method docstring for detailed workflow"
            }
            
        except Exception as e:
            logger.error(f"Error processing query: {e}")
            return {
                "answer": f"I encountered an error while processing your question: {str(e)}",
                "confidence": 0.0,
                "selected_groups": [],
                "search_results": [],
                "metadata_query": False,
                "workflow": "Error occurred during processing"
            }

class AgenticRAGWithLangGraph:
    """Main Agentic RAG system using LangGraph and grouped vector stores"""
    
    def __init__(self, azure_endpoint: str, azure_key: str, api_version: str = "2024-02-15-preview"):
        self.azure_client = AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=azure_key,
            api_version=api_version
        )
        
        self.document_processor = DocumentProcessor(self.azure_client)
        self.vector_store = GroupedVectorStore()
        self.file_analyzer = None
        self.agent = None
        
        logger.info("Initialized AgenticRAG with LangGraph system")
    
    def ingest_folder(self, folder_path: str, save_path: str = "./grouped_vector_store"):
        """Ingest all documents from a folder with group classification"""
        logger.info(f"Starting grouped ingestion of folder: {folder_path}")
        
        if not os.path.exists(folder_path):
            logger.error(f"Folder does not exist: {folder_path}")
            return 0
        
        # Analyze file system
        self.file_analyzer = FileSystemAnalyzer(folder_path)
        self.file_analyzer.set_document_processor(self.document_processor)
        self.file_analyzer.analyze()
        
        # Process all supported files
        all_chunks = []
        processed_files = 0
        group_stats = {group: 0 for group in FILE_GROUPS.keys()}
        
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_ext = Path(file).suffix.lower()
                
                try:
                    if file_ext == '.pdf':
                        chunks = self.document_processor.process_pdf(file_path)
                    elif file_ext == '.docx':
                        chunks = self.document_processor.process_docx(file_path)
                    elif file_ext == '.txt':
                        chunks = self.document_processor.process_txt(file_path)
                    else:
                        continue  # Skip unsupported files
                    
                    if chunks:
                        all_chunks.extend(chunks)
                        processed_files += 1
                        # Count chunks by group
                        for chunk in chunks:
                            group_stats[chunk.group] += 1
                        logger.info(f"Processed {file_path}: {len(chunks)} chunks in group {chunks[0].group}")
                    else:
                        logger.warning(f"No chunks created from {file_path}")
                        
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    continue
        
        if not all_chunks:
            logger.error("No chunks were created from any files")
            return 0
        
        # Add to grouped vector store
        self.vector_store.add_documents(all_chunks)
        
        # Save all groups
        self.vector_store.save_all_groups(save_path)
        
        # Initialize agent
        self.agent = LangGraphAgent(self.azure_client, self.vector_store, self.file_analyzer)
        
        logger.info(f"Ingestion complete. Processed {processed_files} files, created {len(all_chunks)} chunks")
        logger.info(f"Group distribution: {group_stats}")
        return len(all_chunks)
    
    def query(self, question: str) -> Dict[str, Any]:
        """Process a query using the LangGraph agentic approach"""
        logger.info(f"Processing query: {question}")
        
        if not self.vector_store.group_indices or not self.agent:
            return {
                "answer": "System not initialized. Please ingest documents first.",
                "confidence": 0.0,
                "selected_groups": [],
                "search_results": []
            }
        
        try:
            # Use the LangGraph agent to process the query
            response = self.agent.process_query(question)
            return response
        except Exception as e:
            logger.error(f"Error processing query: {e}")
            return {
                "answer": f"I apologize, but I encountered an error while processing your question: {str(e)}",
                "confidence": 0.0,
                "selected_groups": [],
                "search_results": []
            }
    
    def load_vector_store(self, save_path: str = "./grouped_vector_store"):
        """Load a previously saved grouped vector store"""
        logger.info(f"Loading grouped vector store from: {save_path}")
        
        loaded_groups = self.vector_store.load_all_groups(save_path)
        
        if loaded_groups:
            # Initialize agent if we have a file analyzer
            if self.file_analyzer:
                self.agent = LangGraphAgent(self.azure_client, self.vector_store, self.file_analyzer)
            else:
                # Create a minimal file analyzer
                self.file_analyzer = FileSystemAnalyzer("./")
                self.file_analyzer.set_document_processor(self.document_processor)
                self.agent = LangGraphAgent(self.azure_client, self.vector_store, self.file_analyzer)
            
            logger.info(f"Grouped vector store loaded successfully. Groups: {loaded_groups}")
            return True
        else:
            logger.error("Failed to load grouped vector store")
            return False
    
    def get_system_stats(self) -> Dict[str, Any]:
        """Enhanced system statistics"""
        stats = {
            "vector_store_initialized": bool(self.vector_store.group_indices),
            "total_groups": len(self.vector_store.group_indices),
            "agent_initialized": self.agent is not None,
            "file_analyzer_available": self.file_analyzer is not None,
            "supported_file_types": list(self.document_processor.supported_extensions),
            "embedding_model": "all-MiniLM-L6-v2"
        }
        
        # Enhanced group statistics
        if self.vector_store.group_indices:
            enhanced_stats = self.vector_store.get_enhanced_group_stats()
            file_summary = self.vector_store.get_file_type_summary()
            
            stats.update({
                "file_type_summary": file_summary,
                "enhanced_group_statistics": enhanced_stats['groups'],
                "file_type_distribution": enhanced_stats['file_type_distribution'],
                "total_chunks": sum(g['chunk_count'] for g in enhanced_stats['groups'].values()),
                "total_unique_files": sum(g['unique_files'] for g in enhanced_stats['groups'].values()),
                "available_groups": list(self.vector_store.group_indices.keys())
            })
            
            # Add processing statistics if available
            if hasattr(self.document_processor, 'processing_stats'):
                processing_stats = self.document_processor.get_processing_stats()
                stats["processing_statistics"] = processing_stats
        
        return stats
    
    def get_group_file_details(self, group_name: str) -> Dict[str, Any]:
        """Get detailed file information for a specific group"""
        return self.vector_store.get_group_files(group_name)

    def batch_query(self, questions: List[str]) -> List[Dict[str, Any]]:
        """Process multiple queries in batch"""
        results = []
        
        for i, question in enumerate(questions):
            logger.info(f"Processing batch query {i+1}/{len(questions)}: {question}")
            try:
                response = self.query(question)
                results.append({
                    "question": question,
                    "answer": response["answer"],
                    "confidence": response["confidence"],
                    "selected_groups": response["selected_groups"],
                    "status": "success"
                })
            except Exception as e:
                logger.error(f"Error in batch query {i+1}: {e}")
                results.append({
                    "question": question,
                    "answer": f"Error: {str(e)}",
                    "confidence": 0.0,
                    "selected_groups": [],
                    "status": "error"
                })
        
        return results
    
    def export_knowledge_base(self, export_path: str):
        """Export the grouped knowledge base to a readable format"""
        try:
            export_data = {
                "metadata": {
                    "export_timestamp": datetime.now().isoformat(),
                    "total_groups": len(self.vector_store.group_indices),
                    "system_stats": self.get_system_stats()
                },
                "groups": {}
            }
            
            # Export each group's documents
            for group_name, chunks in self.vector_store.group_chunks.items():
                if not chunks:
                    continue
                
                # Group chunks by file
                files_dict = {}
                for chunk in chunks:
                    file_name = chunk.metadata.get('file_name', 'unknown')
                    if file_name not in files_dict:
                        files_dict[file_name] = []
                    files_dict[file_name].append(chunk)
                
                # Export each file's content
                group_data = {
                    "group_name": group_name,
                    "total_chunks": len(chunks),
                    "total_files": len(files_dict),
                    "documents": []
                }
                
                for file_name, file_chunks in files_dict.items():
                    file_data = {
                        "file_name": file_name,
                        "total_chunks": len(file_chunks),
                        "file_type": file_chunks[0].metadata.get('file_type', 'unknown'),
                        "processed_at": file_chunks[0].metadata.get('processed_at', 'unknown'),
                        "chunks": [
                            {
                                "chunk_id": chunk.chunk_id,
                                "content": chunk.content,
                                "metadata": chunk.metadata
                            } for chunk in file_chunks
                        ]
                    }
                    group_data["documents"].append(file_data)
                
                export_data["groups"][group_name] = group_data
            
            # Save to JSON
            with open(export_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Grouped knowledge base exported to {export_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error exporting knowledge base: {e}")
            return False

def enhanced_sidebar_display(st, rag_system):
    """Enhanced sidebar with professional styling"""
    
    # Custom sidebar styling
    st.markdown("""
    <style>
    .sidebar .sidebar-content {
        background-color: #2c3e50;
        color: white;
    }
    .sidebar .sidebar-content .stRadio > div > label {
        color: white;
    }
    .sidebar .sidebar-content .stTextInput > div > div > input {
        background-color: #34495e;
        color: white;
    }
    .sidebar .sidebar-content .stButton > button {
        background-color: #3498db;
        color: white;
        border: none;
        border-radius: 4px;
        padding: 8px 16px;
    }
    .sidebar .sidebar-content .stButton > button:hover {
        background-color: #2980b9;
    }
    .sidebar .sidebar-content .stExpander > div > div {
        background-color: #34495e;
        border-radius: 4px;
        padding: 8px;
        margin-bottom: 8px;
    }
    </style>
    """, unsafe_allow_html=True)
    
    with st.sidebar:
        # Sidebar header with accent color
        st.markdown("""
        <div style="background-color: #3498db; padding: 10px; border-radius: 4px; margin-bottom: 20px;">
            <h2 style="color: white; margin: 0;">Configuration</h2>
        </div>
        """, unsafe_allow_html=True)
        
        # Action selection with custom styling
        action = st.radio("Choose Action", ["Ingest New Folder", "Load Vector Store"])
        
        if action == "Ingest New Folder":
            folder_path = st.text_input("Enter folder path to ingest:", "./Input")
            if st.button("Ingest Documents", key="ingest_btn") and folder_path:
                if os.path.exists(folder_path):
                    with st.spinner(f"Ingesting documents from {folder_path}..."):
                        chunks_created = rag_system.ingest_folder(folder_path)
                        st.success(f"Created {chunks_created} chunks from documents")
                        st.rerun()
                else:
                    st.error(f"Folder {folder_path} does not exist.")
        
        elif action == "Load Vector Store":
            vector_store_path = st.text_input("Enter vector store path:", "./vector_store")
            if st.button("Load Vector Store", key="load_btn"):
                with st.spinner("Loading vector store..."):
                    if rag_system.load_vector_store(vector_store_path):
                        st.success("Vector store loaded successfully")
                        st.rerun()
                    else:
                        st.error("Failed to load vector store")
        
        # Divider with custom style
        st.markdown("""
        <div style="height: 1px; background-color: #3498db; margin: 20px 0;"></div>
        """, unsafe_allow_html=True)
        
        # Save section with custom styling
        st.markdown("""
        <div style="background-color: #34495e; padding: 10px; border-radius: 4px; margin-bottom: 20px;">
            <h3 style="color: white; margin: 0;">Save Vector Store</h3>
        </div>
        """, unsafe_allow_html=True)
        
        save_path = st.text_input("Save path:", "./grouped_vector_store")
        if st.button("Save Vector Store", key="save_btn"):
            if rag_system.save_vector_store(save_path):
                st.success("Vector store saved successfully")
            else:
                st.error("Failed to save vector store")
        
        # Document groups section
        st.markdown("""
        <div style="background-color: #34495e; padding: 10px; border-radius: 4px; margin: 20px 0;">
            <h3 style="color: white; margin: 0;">Document Groups</h3>
        </div>
        """, unsafe_allow_html=True)
        
        stats = rag_system.get_system_stats()
        
        for group_name, keywords in FILE_GROUPS.items():
            group_stats = stats.get('enhanced_group_statistics', {}).get(group_name, {})
            file_count = group_stats.get('unique_files', 0)
            
            with st.expander(f"{group_name} ({file_count} files)"):
                if keywords:
                    st.markdown(f"<p style='color: #3498db;'><strong>Keywords:</strong> {', '.join(keywords)}</p>", unsafe_allow_html=True)
                else:
                    st.markdown("<p style='color: #3498db;'><strong>Default group</strong> for unmatched files</p>", unsafe_allow_html=True)
                
                if file_count > 0:
                    st.markdown(f"<p><strong>Total chunks:</strong> {group_stats.get('chunk_count', 0)}</p>", unsafe_allow_html=True)
                    
                    file_details = rag_system.get_group_file_details(group_name)
                    
                    st.markdown("<p><strong>Files in this group:</strong></p>", unsafe_allow_html=True)
                    for file_name in file_details.get('files', []):
                        file_info = file_details.get('file_details', {}).get(file_name, {})
                        processing_method = file_info.get('processing_method', 'Unknown')
                        chunks = file_info.get('chunks', 0)
                        
                        st.markdown(f"<p style='margin-left: 20px;'>{file_name} ({chunks} chunks)</p>", unsafe_allow_html=True)
                        
                        if file_info.get('total_pages', 0) > 0:
                            vision_pages = len(file_info.get('vision_pages', []))
                            text_pages = len(file_info.get('text_pages', []))
                            st.markdown(f"<p style='margin-left: 20px;'>Pages: {text_pages} text, {vision_pages} OCR</p>", unsafe_allow_html=True)
                
                else:
                    st.markdown("<p style='color: #7f8c8d;'>No files in this group</p>", unsafe_allow_html=True)

def enhanced_main_display(st, rag_system):
    """Enhanced main display with professional styling"""
    
    # Custom main page styling
    st.markdown("""
    <style>
    .main .block-container {
        background-color: #f5f7fa;
        padding: 2rem;
    }
    .stMetric {
        background-color: white;
        border-radius: 8px;
        padding: 15px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .stMetric > div > div {
        color: #2c3e50;
    }
    .stMetric > div > div:first-child {
        font-size: 1rem;
        color: #7f8c8d;
    }
    .stMetric > div > div:nth-child(2) {
        font-size: 1.5rem;
        color: #2c3e50;
        font-weight: bold;
    }
    .stDataFrame {
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .stButton > button {
        background-color: #3498db;
        color: white;
        border: none;
        border-radius: 4px;
        padding: 8px 16px;
    }
    .stButton > button:hover {
        background-color: #2980b9;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Main header with gradient background
    st.markdown("""
    <div style="background: linear-gradient(135deg, #3498db, #2c3e50); 
                padding: 20px; 
                border-radius: 8px; 
                margin-bottom: 30px;
                color: white;">
        <h1 style="color: white; margin: 0;">Agentic RAG System</h1>
        <p style="margin: 5px 0 0 0;">Document processing with intelligent group selection and enhanced metadata tracking</p>
    </div>
    """, unsafe_allow_html=True)
    
    # System stats display
    stats = rag_system.get_system_stats()
    
    st.markdown("""
    <div style="background-color: white; 
                padding: 15px; 
                border-radius: 8px; 
                margin-bottom: 20px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
        <h2 style="color: #2c3e50; margin: 0 0 15px 0;">System Statistics</h2>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Vector Store", "✅ Ready" if stats['vector_store_initialized'] else "❌ Not Ready")
    with col2:
        st.metric("Total Groups", stats['total_groups'])
    with col3:
        st.metric("Total Chunks", stats.get('total_chunks', 0))
    with col4:
        st.metric("Unique Files", stats.get('total_unique_files', 0))
    
    # File type and processing method breakdown
    if stats.get('file_type_summary'):
        st.markdown("""
        <div style="background-color: white; 
                    padding: 15px; 
                    border-radius: 8px; 
                    margin: 20px 0;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
            <h2 style="color: #2c3e50; margin: 0 0 15px 0;">File Processing Summary</h2>
        </div>
        """, unsafe_allow_html=True)
        
        file_summary = stats['file_type_summary']
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("PDF Files", file_summary['pdf_files'])
        with col2:
            st.metric("DOCX Files", file_summary['docx_files'])
        with col3:
            st.metric("TXT Files", file_summary['txt_files'])
        
        # Processing methods
        st.markdown("""
        <div style="background-color: white; 
                    padding: 15px; 
                    border-radius: 8px; 
                    margin: 20px 0;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
            <h2 style="color: #2c3e50; margin: 0 0 15px 0;">Processing Methods</h2>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Text Extraction", file_summary['processing_methods']['text_extraction'])
        with col2:
            st.metric("Vision OCR", file_summary['processing_methods']['vision_ocr'])
    
    # Group statistics table
    if stats.get('enhanced_group_statistics'):
        st.markdown("""
        <div style="background-color: white; 
                    padding: 15px; 
                    border-radius: 8px; 
                    margin: 20px 0;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
            <h2 style="color: #2c3e50; margin: 0 0 15px 0;">Detailed Group Statistics</h2>
        </div>
        """, unsafe_allow_html=True)

class ChatBot:
    """Interactive chatbot for the Agentic RAG system"""
    
    def __init__(self, rag_system: AgenticRAGWithLangGraph):
        self.rag_system = rag_system
        self.conversation_history = []
    
    def run(self):
        """Run the interactive chatbot"""
        print("🤖 Welcome to the Agentic RAG System!")
        print("Type 'exit' or 'quit' to end the conversation.")
        print("Type 'stats' to see system statistics.")
        print("Type 'groups' to see available document groups.")
        print("-" * 50)
        
        while True:
            try:
                user_input = input("\n💬 You: ").strip()
                
                if user_input.lower() in ['exit', 'quit']:
                    print("👋 Goodbye!")
                    break
                
                if user_input.lower() == 'stats':
                    stats = self.rag_system.get_system_stats()
                    self._display_stats(stats)
                    continue
                
                if user_input.lower() == 'groups':
                    self._display_groups()
                    continue
                
                if not user_input:
                    continue
                
                print("\n🔍 Processing your query...")
                response = self.rag_system.query(user_input)
                
                print(f"\n🤖 Assistant: {response['answer']}")
                print(f"📊 Confidence: {response['confidence']:.2f}")
                print(f"📂 Searched Groups: {', '.join(response['selected_groups'])}")
                
                # Store conversation
                self.conversation_history.append({
                    "user": user_input,
                    "assistant": response['answer'],
                    "confidence": response['confidence'],
                    "groups": response['selected_groups']
                })
                
            except KeyboardInterrupt:
                print("\n👋 Goodbye!")
                break
            except Exception as e:
                print(f"\n❌ Error: {e}")
    
    def _display_stats(self, stats: Dict[str, Any]):
        """Display system statistics"""
        print("\n📈 System Statistics:")
        print(f"  Vector Store Initialized: {'✅' if stats['vector_store_initialized'] else '❌'}")
        print(f"  Total Groups: {stats['total_groups']}")
        print(f"  Total Chunks: {stats.get('total_chunks', 0)}")
        print(f"  Total Files: {stats.get('total_unique_files', 0)}")
        print(f"  Available Groups: {', '.join(stats.get('available_groups', []))}")
        
        if stats.get('enhanced_group_statistics'):
            print("\n📊 Group Statistics:")
            for group, group_stats in stats['enhanced_group_statistics'].items():
                if group_stats['unique_files'] > 0:
                    print(f"  {group}: {group_stats['chunk_count']} chunks, {group_stats['unique_files']} files")
    
    def _display_groups(self):
        """Display available document groups"""
        print("\n📂 Available Document Groups:")
        for group_name, keywords in FILE_GROUPS.items():
            if keywords:
                print(f"  {group_name}: {', '.join(keywords)}")
            else:
                print(f"  {group_name}: Default group for unmatched files")

def main():
    """Main function to run the enhanced Streamlit demo"""
    st.set_page_config(
        page_title="Agentic RAG System", 
        layout="wide"
    )
    

    # Configuration - Load from environment variables with fallback to defaults
 
    
    AZURE_ENDPOINT = "https://oai-nasco.openai.azure.com/openai/deployments/gpt-4o/chat/completions?api-version=2025-01-01-preview"
    AZURE_KEY = "YOUR_AZURE_OPENAI_API_KEY"
    API_VERSION = "2024-02-15-preview"
        
    

    
    # Initialize system if not already in session state
    if 'rag_system' not in st.session_state:
        st.session_state.rag_system = AgenticRAGWithLangGraph(AZURE_ENDPOINT, AZURE_KEY, API_VERSION)
    
    rag_system = st.session_state.rag_system
    
    # Enhanced sidebar display
    enhanced_sidebar_display(st, rag_system)
    
    # Enhanced main display
    enhanced_main_display(st, rag_system)
    
    

    # Query Interface
    st.subheader("Ask a Question")
    
    question = st.text_input(
        "Enter your question about the documents:", 
        placeholder="e.g., How many PDF files are there?",
        key="question_input"
    )
    
    # Create button columns
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Submit Query", key="submit_query") and question:
            with st.spinner("Processing your query..."):
                response = rag_system.query(question)
                
                # Display results
                st.markdown("### Answer:")
                st.write(response['answer'])
                
                # Display metadata
                col_meta1, col_meta2 = st.columns(2)
                with col_meta1:
                    st.metric("Confidence", f"{response['confidence']:.2f}")
                with col_meta2:
                    st.write(f"**Groups Searched:** {', '.join(response['selected_groups'])}")
                    
                    # Display search results if available
                    if response.get('search_results'):
                        with st.expander("Search Results Details"):
                            for i, result in enumerate(response['search_results'][:3]):
                                st.markdown(f"**Result {i+1}** (Score: {result['score']:.3f})")
                                st.markdown(f"Group: {result['group']}")
                                st.markdown(f"File: {result['file_name']}")
                                st.markdown(f"Content: {result['content'][:200]}...")
                                st.markdown("---")
                                
    
if __name__ == "__main__":
    main()
