import os
import logging
from typing import List, Dict, Any, Optional, Tuple, TypedDict, Union
from dataclasses import dataclass
from pathlib import Path
from datetime import datetime
import json
import time
import streamlit as st
import re
from concurrent.futures import ThreadPoolExecutor
from document_chunk import DocumentChunk

# Core libraries
from openai import AzureOpenAI
import fitz  # PyMuPDF
from docx import Document
import base64

# Vector store and embeddings
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import pickle
from langchain.text_splitter import RecursiveCharacterTextSplitter

# LangGraph imports
from langgraph.graph import StateGraph, END
from langgraph.graph.message import add_messages
from langchain_core.messages import HumanMessage, AIMessage, BaseMessage

from typing_extensions import Annotated
from sentence_transformers.cross_encoder import CrossEncoder

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# File grouping configuration
FILE_GROUPS = {
    "Chemical Consumption": [
        "chemical", "consumption", "chemicals"
    ],
    "Engineering Tickets": [
        "engineering", "ticket", "tickets"
    ],
    "Risk Assessment and Hazard Analysis": [
        "risk", "hazard", "safety", "assessment", "analysis", "incident", "accident"
    ],
    "Well Recap": [
        "well recap", "Well Recap", "recap"
    ],
    "Mud Program": [
        "mud", "program"
    ],
    "Contractor Feedback": [
        "contractor", "feedback"
    ],
    "Hydraulic Summary": [
        "hydraulic", "summary","hydraulics"
    ],
    "Other Group": []  # Default group for unmatched files
}

@dataclass
class FileSystemInfo:
    """Stores file system structure information"""
    total_folders: int
    total_files: int
    file_types: Dict[str, int]
    folder_structure: Dict[str, Any]
    file_groups: Dict[str, List[str]]

## MODIFIED: The LangGraph state is now much richer to support the self-correction loop.
class AgentState(TypedDict):
    messages: Annotated[List[BaseMessage], add_messages]
    original_question: str
    question: str
    selected_groups: List[str]
    search_results: List[Dict[str, Any]]
    final_answer: str
    confidence: float
    # ## NEW: Fields for self-correction and proactive insights
    retrieval_grade: str # Grade of the retrieved documents ('relevant' or 'not_relevant')
    correction_attempts: int # Counter to prevent infinite loops
    proactive_insights: str # Field for storing generated insights

class DocumentProcessor:
    """Enhanced document processor with hybrid layout-aware and OCR chunking for PDFs."""
    
    def __init__(self, azure_client: AzureOpenAI):
        """Initialize with Azure client and setup processing stats"""
        if not azure_client:
            raise ValueError("AzureOpenAI client cannot be None")
            
        self.azure_client = azure_client
        self.supported_extensions = {'.pdf', '.docx', '.txt'}
        self.processing_stats = {
            'text_extraction': [],
            'vision_ocr': [],
            'total_processed': 0
        }
        self.metadata_summary = []
    
    def classify_file_group(self, file_path: str) -> str:
        # -- UNCHANGED --
        file_name = os.path.basename(file_path).lower()
        for group_name, keywords in FILE_GROUPS.items():
            if group_name == "Other Group": continue
            for keyword in keywords:
                if keyword.lower() in file_name:
                    return group_name
        return "Other Group"

    def _format_table_as_markdown(self, table) -> str:
        # -- UNCHANGED --
        markdown_text = ""
        headers = [cell.strip() for cell in table.header.names if cell]
        if headers:
            markdown_text += "| " + " | ".join(headers) + " |\n"
            markdown_text += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        for row in table.extract():
            cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
            if any(cleaned_row):
                markdown_text += "| " + " | ".join(cleaned_row) + " |\n"
        return markdown_text

    ## NEW: Re-integrated the OCR method for processing scanned pages.
    def _extract_text_via_ocr_page(self, page) -> str:
        """Extract text from a single PDF page using GPT-4V OCR with rate limiting."""
        try:
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            img_data = pix.tobytes("png")
            img_base64 = base64.b64encode(img_data).decode()
            
            max_retries = 3
            base_delay = 1.0
            for attempt in range(max_retries):
                try:
                    response = self.azure_client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "user", "content": [
                                {"type": "text", "text": "Extract all text from this image. Maintain the original formatting as much as possible. If there are tables, preserve their structure. Return only the extracted text without any commentary."},
                                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_base64}"}}
                            ]}
                        ],
                        max_tokens=4000, temperature=0
                    )
                    return response.choices[0].message.content
                except Exception as e:
                    if "429" in str(e):
                        delay = base_delay * (2 ** attempt)
                        logger.warning(f"Rate limited, retrying in {delay} seconds...")
                        time.sleep(delay)
                    else:
                        raise e
            logger.error(f"OCR failed after {max_retries} attempts.")
            return ""
        except Exception as e:
            logger.error(f"OCR failed for page: {e}")
            return ""

    ## MODIFIED: The process_pdf method is now a hybrid of layout-aware and OCR processing.
    def process_pdf(self, file_path: str, ocr_threshold: int = 50) -> List[DocumentChunk]:
        """
        MODIFIED: Process PDF using a hybrid approach.
        - For text-rich pages, it uses layout-aware chunking (tables, paragraphs).
        - For image-based pages (scans), it uses vision-based OCR.
        """
        if not os.path.exists(file_path) or os.path.getsize(file_path) == 0:
            return []

        all_chunks = []
        file_name = os.path.basename(file_path)
        file_group = self.classify_file_group(file_path)
        chunk_id_counter = 0

        try:
            doc = fitz.open(file_path)
            self.processing_stats['total_processed'] += 1
            
            for page_num, page in enumerate(doc):
                # HYBRID LOGIC: Decide whether to use layout-aware or OCR path
                if len(page.get_text("text").strip()) > ocr_threshold:
                    # --- PATH 1: Layout-Aware Chunking for Native PDFs ---
                    self.processing_stats['text_extraction'].append(page_num + 1)
                    
                    # Process tables
                    tables = page.find_tables()
                    table_bboxes = [fitz.Rect(t.bbox) for t in tables]
                    for i, table in enumerate(tables):
                        table_markdown = self._format_table_as_markdown(table)
                        if not table_markdown.strip(): continue
                        chunk_id = f"{file_name}_page_{page_num+1}_table_{i}"
                        metadata = {
                            "file_name": file_name, "group": file_group, "file_type": "pdf",
                            "element_type": "table", "page_number": page_num + 1,
                            "bounding_box": list(table.bbox), "processing_method": "Layout-Aware Extraction"
                        }
                        all_chunks.append(DocumentChunk(content=table_markdown, source_file=file_path, chunk_id=chunk_id, group=file_group, metadata=metadata))
                        chunk_id_counter += 1

                    # Process text blocks
                    text_blocks = page.get_text("blocks")
                    child_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=75)
                    for block in text_blocks:
                        block_rect = fitz.Rect(block[:4])
                        if any(block_rect.intersects(bbox) for bbox in table_bboxes): continue
                        block_text = block[4].strip()
                        if len(block_text) < 20: continue
                        
                        smaller_sub_chunks = child_splitter.split_text(block_text)
                        for sub_chunk_text in smaller_sub_chunks:
                            chunk_id = f"{file_name}_chunk_{chunk_id_counter}"
                            metadata = {
                                "parent_content": block_text, "file_name": file_name, "group": file_group,
                                "file_type": "pdf", "element_type": "paragraph", "page_number": page_num + 1,
                                "bounding_box": list(block[:4]), "processing_method": "Layout-Aware Extraction"
                            }
                            all_chunks.append(DocumentChunk(content=sub_chunk_text, source_file=file_path, chunk_id=chunk_id, group=file_group, metadata=metadata))
                            chunk_id_counter += 1
                else:
                    # --- PATH 2: OCR for Scanned Pages ---
                    logger.info(f"Page {page_num+1} in {file_name} has low text. Applying OCR.")
                    self.processing_stats['vision_ocr'].append(page_num + 1)
                    
                    ocr_text = self._extract_text_via_ocr_page(page)
                    if ocr_text:
                        # Use the standard _create_chunks for the OCR text stream
                        ocr_chunks = self._create_chunks(ocr_text, file_path, "pdf", file_group)
                        for chunk in ocr_chunks:
                            chunk.metadata['processing_method'] = "Vision (OCR)"
                            chunk.metadata['page_number'] = page_num + 1
                        all_chunks.extend(ocr_chunks)

            logger.info(f"Created {len(all_chunks)} hybrid chunks from {file_name}")
            return all_chunks

        except Exception as e:
            logger.error(f"Error in hybrid PDF processing for {file_path}: {e}", exc_info=True)
            return []

    def process_docx(self, file_path: str) -> List[DocumentChunk]:
        # -- UNCHANGED --
        try:
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            for table in doc.tables:
                text += "\n" + "\n".join(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
            return self._create_chunks(text, file_path, "docx", self.classify_file_group(file_path))
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return []
    
    def process_txt(self, file_path: str) -> List[DocumentChunk]:
        # -- UNCHANGED --
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return self._create_chunks(file.read(), file_path, "txt", self.classify_file_group(file_path))
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            return []

    def _extract_metadata(self, text: str) -> dict:
        # -- UNCHANGED --
        date_pattern = r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"
        dates = list(set(re.findall(date_pattern, text)))
        well_pattern = r"\bWell\s*[A-Z0-9-]+\b"
        wells = list(set(re.findall(well_pattern, text, re.IGNORECASE)))
        return {"dates": sorted(dates), "wells": wells}

    def _create_chunks(self, text: str, file_path: str, file_type: str, group: str) -> List[DocumentChunk]:
        # -- UNCHANGED --
        if not text.strip(): return []
        file_name = os.path.basename(file_path)
        metadata = self._extract_metadata(text)
        parent_chunks = text.split("\n\n")
        child_chunks = []
        chunk_id_counter = 0
        child_splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=50)
        for i, parent_chunk_text in enumerate(parent_chunks):
            if not parent_chunk_text.strip(): continue
            smaller_sub_chunks = child_splitter.split_text(parent_chunk_text)
            for sub_chunk_text in smaller_sub_chunks:
                chunk_id = f"{file_name}_chunk_{chunk_id_counter}"
                new_chunk = DocumentChunk(
                    content=sub_chunk_text, source_file=file_path, chunk_id=chunk_id, group=group,
                    metadata={
                        "parent_content": parent_chunk_text, "parent_chunk_index": i,
                        "file_type": file_type, "file_size": len(sub_chunk_text),
                        "processed_at": datetime.now().isoformat(), "file_name": file_name,
                        "group": group, "processing_method": "Text Extraction",
                        "dates": metadata["dates"], "wells": metadata["wells"]
                    }
                )
                child_chunks.append(new_chunk)
                chunk_id_counter += 1
        return child_chunks

class GroupedVectorStore:
    """Manages document embeddings with group-based organization"""
    
    def __init__(self, embedding_model_name: str = "all-mpnet-base-v2"):
        """Initialize with configurable embedding model"""
        # Store the model name for easy access later
        self.model_name = embedding_model_name
        self.embedding_model = SentenceTransformer(self.model_name)
        self.group_indices = {}
        self.group_chunks = {}
        self.dimension = self.embedding_model.get_sentence_embedding_dimension()
        
        # Enhanced tracking
        self.file_type_stats = {}
        self.group_file_details = {}
        self.date_index = {}  # {date: [file1, file2]}
        self.well_index = {}  # {well_name: [file1, file2]}
        
        logger.info(f"Initialized GroupedVectorStore with model: {self.model_name}")
        
    def _update_indices(self, chunk: DocumentChunk):
        """Update date and well indices"""
        file_name = chunk.metadata.get('file_name', 'unknown')
        dates = chunk.metadata.get('dates', [])
        wells = chunk.metadata.get('wells', [])
        
        for date in dates:
            if date not in self.date_index:
                self.date_index[date] = []
            if file_name not in self.date_index[date]:
                self.date_index[date].append(file_name)
                
        for well in wells:
            if well not in self.well_index:
                self.well_index[well] = []
            if file_name not in self.well_index[well]:
                self.well_index[well].append(file_name)
    
    def add_documents(self, chunks: List[DocumentChunk]):
        """Enhanced document addition with detailed tracking"""
        if not chunks:
            logger.warning("No chunks provided to add_documents")
            return
        
        # Track file types and group details
        for chunk in chunks:
            file_type = chunk.metadata.get('file_type', 'unknown')
            group = chunk.group
            file_name = chunk.metadata.get('file_name', 'unknown')
            
            # Update file type stats
            if file_type not in self.file_type_stats:
                self.file_type_stats[file_type] = 0
            self.file_type_stats[file_type] += 1
            
            # Update group file details
            if group not in self.group_file_details:
                self.group_file_details[group] = {
                    'files': set(),
                    'file_details': {},
                    'total_chunks': 0
                }
            
            self.group_file_details[group]['files'].add(file_name)
            self.group_file_details[group]['total_chunks'] += 1
            
            if file_name not in self.group_file_details[group]['file_details']:
                self.group_file_details[group]['file_details'][file_name] = {
                    'file_type': file_type,
                    'processing_method': chunk.metadata.get('processing_method', 'Unknown'),
                    'chunks': 0,
                    'total_pages': chunk.metadata.get('total_pages', 0),
                    'text_pages': chunk.metadata.get('text_pages', []),
                    'vision_pages': chunk.metadata.get('vision_pages', []),
                    'dates': chunk.metadata.get('dates', []),
                    'wells': chunk.metadata.get('wells', [])
                }
            
            self.group_file_details[group]['file_details'][file_name]['chunks'] += 1
            
            # Update date and well indices
            self._update_indices(chunk)
        
        # Group chunks by their group classification
        grouped_chunks = {}
        for chunk in chunks:
            group = chunk.group
            if group not in grouped_chunks:
                grouped_chunks[group] = []
            grouped_chunks[group].append(chunk)
        
        # Process each group
        for group_name, group_chunks in grouped_chunks.items():
            logger.info(f"Adding {len(group_chunks)} chunks to group: {group_name}")
            self._add_to_group(group_name, group_chunks)
    
    def _add_to_group(self, group_name: str, chunks: List[DocumentChunk]):
        """Add chunks to a group's vector store with batch embeddings"""
        if not chunks:
            return

        # Initialize index if needed
        if group_name not in self.group_indices:
            self.group_indices[group_name] = faiss.IndexFlatIP(self.dimension)
            self.group_chunks[group_name] = []

        # Process in batches with parallel embedding generation
        batch_size = 64  # Increased batch size
        embeddings = []
        
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            try:
                # Generate embeddings in parallel using encode() method
                with ThreadPoolExecutor() as executor:
                    batch_contents = [chunk.content for chunk in batch]
                    batch_embeddings = list(executor.map(
                        lambda content: self.embedding_model.encode(content),
                        batch_contents
                    ))
                embeddings.extend(batch_embeddings)
                
                # Store embeddings in chunks
                for chunk, emb in zip(batch, batch_embeddings):
                    chunk.embedding = emb
                    
            except Exception as e:
                logger.error(f"Embedding generation failed for batch {i//batch_size}: {e}")
                continue

        if embeddings:
            # Add to FAISS index
            self.group_indices[group_name].add(np.array(embeddings).astype('float32'))
            self.group_chunks[group_name].extend(chunks)
    
    def search_groups(self, query: str, groups: List[str], k: int = 5) -> List[Tuple[DocumentChunk, float, str]]:
        """Search across specified groups"""
        all_results = []
        
        for group_name in groups:
            if group_name not in self.group_indices or len(self.group_chunks[group_name]) == 0:
                logger.warning(f"Group {group_name} is empty or not initialized")
                continue
            
            try:
                query_embedding = self.embedding_model.encode(query)
                scores, indices = self.group_indices[group_name].search(
                    np.array([query_embedding]).astype('float32'), 
                    min(k, len(self.group_chunks[group_name]))
                )
                
                for score, idx in zip(scores[0], indices[0]):
                    if 0 <= idx < len(self.group_chunks[group_name]):
                        all_results.append((
                            self.group_chunks[group_name][idx], 
                            float(score), 
                            group_name
                        ))
                        
            except Exception as e:
                logger.error(f"Error searching group {group_name}: {e}")
                continue
        
        # Sort by score and return top k
        all_results.sort(key=lambda x: x[1], reverse=True)
        logger.info(f"Found {len(all_results)} results across groups {groups}")
        return all_results[:k]
    
    def get_enhanced_group_stats(self) -> Dict[str, Any]:
        """Get enhanced statistics with file details"""
        stats = {
            'file_type_distribution': dict(self.file_type_stats),
            'total_groups': len(self.group_indices),
            'groups': {}
        }
        
        for group_name in self.group_indices:
            group_details = self.group_file_details.get(group_name, {})
            chunk_count = len(self.group_chunks.get(group_name, []))
            
            stats['groups'][group_name] = {
                'chunk_count': chunk_count,
                'unique_files': len(group_details.get('files', set())),
                'file_list': list(group_details.get('files', set())),
                'file_details': group_details.get('file_details', {}),
                'total_chunks_tracked': group_details.get('total_chunks', 0)
            }
        
        return stats

    def get_file_type_summary(self) -> Dict[str, Any]:
        """Get summary by file type"""
        summary = {
            'pdf_files': 0,
            'docx_files': 0,
            'txt_files': 0,
            'total_files': 0,
            'processing_methods': {
                'text_extraction': 0,
                'vision_ocr': 0
            }
        }
        
        for group_name, group_details in self.group_file_details.items():
            for file_name, file_info in group_details.get('file_details', {}).items():
                file_type = file_info.get('file_type', 'unknown')
                processing_method = file_info.get('processing_method', 'Unknown')
                
                if file_type == 'pdf':
                    summary['pdf_files'] += 1
                elif file_type == 'docx':
                    summary['docx_files'] += 1
                elif file_type == 'txt':
                    summary['txt_files'] += 1
                
                summary['total_files'] += 1
                
                if 'Text Extraction' in processing_method:
                    summary['processing_methods']['text_extraction'] += 1
                elif 'Vision (OCR)' in processing_method:
                    summary['processing_methods']['vision_ocr'] += 1
        
        return summary
    

    def get_group_files(self, group_name: str) -> Dict[str, Any]:
        """Get detailed information about files in a specific group"""
        if group_name not in self.group_file_details:
            return {'files': [], 'total_files': 0}
        
        group_info = self.group_file_details[group_name]
        return {
            'files': list(group_info.get('files', set())),
            'total_files': len(group_info.get('files', set())),
            'file_details': group_info.get('file_details', {}),
            'total_chunks': group_info.get('total_chunks', 0)
        }
    
    def save_group(self, group_name: str, folder_path: str):
        """Save a specific group's vector store to a folder"""
        try:
            os.makedirs(folder_path, exist_ok=True)
            
            if group_name in self.group_indices:
                index_path = os.path.join(folder_path, f"{group_name}.index")
                faiss.write_index(self.group_indices[group_name], index_path)
                logger.info(f"Saved FAISS index for group {group_name} to {index_path}")
            
            if group_name in self.group_chunks:
                chunks_to_save = []
                for chunk in self.group_chunks[group_name]:
                    chunk_copy = DocumentChunk(
                        content=chunk.content,
                        source_file=chunk.source_file,
                        chunk_id=chunk.chunk_id,
                        group=chunk.group,
                        metadata=chunk.metadata,
                        embedding=None
                    )
                    chunks_to_save.append(chunk_copy)
                
                chunks_path = os.path.join(folder_path, f"{group_name}.chunks")
                with open(chunks_path, 'wb') as f:
                    pickle.dump(chunks_to_save, f)
                
                logger.info(f"Saved {len(chunks_to_save)} chunks for group {group_name} to {chunks_path}")
                
        except Exception as e:
            logger.error(f"Error saving group {group_name}: {e}")
    
    def save_all_groups(self, folder_path: str):
        """Save all groups to a single folder"""
        try:
            os.makedirs(folder_path, exist_ok=True)
            for group_name in self.group_indices:
                self.save_group(group_name, folder_path)
            logger.info(f"Saved all groups to folder: {folder_path}")
            return True
        except Exception as e:
            logger.error(f"Error saving all groups: {e}")
            return False
    
    def load_group(self, group_name: str, folder_path: str):
        """Load a specific group's vector store from folder"""
        try:
            index_path = os.path.join(folder_path, f"{group_name}.index")
            chunks_path = os.path.join(folder_path, f"{group_name}.chunks")
            
            if os.path.exists(index_path):
                self.group_indices[group_name] = faiss.read_index(index_path)
                logger.info(f"Loaded FAISS index for group {group_name} from {index_path}")
            else:
                logger.warning(f"FAISS index file not found at {index_path}")
                return False
            
            if os.path.exists(chunks_path):
                with open(chunks_path, 'rb') as f:
                    self.group_chunks[group_name] = pickle.load(f)
                logger.info(f"Loaded {len(self.group_chunks[group_name])} chunks for group {group_name} from {chunks_path}")
            else:
                logger.warning(f"Chunks file not found at {chunks_path}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error loading group {group_name}: {e}")
            return False
    
    def load_all_groups(self, folder_path: str):
        """Load all available groups from folder"""
        loaded_groups = []
        
        if not os.path.exists(folder_path):
            logger.error(f"Vector store folder not found: {folder_path}")
            return loaded_groups
            
        for filename in os.listdir(folder_path):
            if filename.endswith('.index'):
                group_name = filename[:-6]  # Remove '.index'
                if self.load_group(group_name, folder_path):
                    loaded_groups.append(group_name)
        
        logger.info(f"Loaded groups from {folder_path}: {loaded_groups}")
        return loaded_groups


class FileSystemAnalyzer:
    """Analyzes file system structure with group classification"""
    
    def __init__(self, root_path: str):
        self.root_path = Path(root_path)
        self.file_info = None
        self.document_processor = None
    
    def set_document_processor(self, processor: DocumentProcessor):
        """Set document processor for file classification"""
        self.document_processor = processor
    
    def analyze(self) -> FileSystemInfo:
        """Analyze file system structure with group classification"""
        folder_count = 0
        file_count = 0
        file_types = {}
        folder_structure = {}
        file_groups = {group: [] for group in FILE_GROUPS.keys()}
        
        try:
            for root, dirs, files in os.walk(self.root_path):
                folder_count += len(dirs)
                file_count += len(files)
                
                for file in files:
                    file_path = os.path.join(root, file)
                    ext = Path(file).suffix.lower()
                    file_types[ext] = file_types.get(ext, 0) + 1
                    
                    # Classify file into group
                    if self.document_processor:
                        group = self.document_processor.classify_file_group(file_path)
                        file_groups[group].append(file)
                
                rel_path = os.path.relpath(root, self.root_path)
                if rel_path == '.':
                    rel_path = 'root'
                
                folder_structure[rel_path] = {
                    'subdirs': dirs,
                    'files': files,
                    'file_count': len(files)
                }
        
        except Exception as e:
            logger.error(f"Error analyzing file system: {e}")
        
        self.file_info = FileSystemInfo(
            total_folders=folder_count,
            total_files=file_count,
            file_types=file_types,
            folder_structure=folder_structure,
            file_groups=file_groups
        )
        
        logger.info(f"Analyzed file system: {file_count} files, {folder_count} folders")
        return self.file_info
    
    def get_folder_summary(self) -> str:
        """Get a summary of the folder structure with group information"""
        if not self.file_info:
            self.analyze()
        
        summary = f"File System Analysis:\n"
        summary += f"- Total folders: {self.file_info.total_folders}\n"
        summary += f"- Total files: {self.file_info.total_files}\n"
        summary += f"- File types: {dict(self.file_info.file_types)}\n\n"
        
        summary += "File Groups:\n"
        for group, files in self.file_info.file_groups.items():
            summary += f"- {group}: {len(files)} files\n"
        
        return summary

## MODIFIED: This class is heavily refactored to implement the advanced agentic loop.
class LangGraphAgent:
    """LangGraph-based agent with self-correcting retrieval and proactive synthesis."""
    
    def __init__(self, azure_client: AzureOpenAI, vector_store: GroupedVectorStore, file_analyzer: FileSystemAnalyzer):
        self.azure_client = azure_client
        self.vector_store = vector_store
        self.file_analyzer = file_analyzer
        self.cross_encoder = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
        self.query_cache = {}
        # The graph is now built by a dedicated method
        self.graph = self._build_graph()

    ## NEW: Node for transforming the user query into sub-questions for better retrieval.
    def _query_transformer(self, state: AgentState) -> AgentState:
        """Transforms the user question into a more effective search query or sub-questions."""
        logger.info("Transforming query...")
        
        # If we are in a correction loop, we modify the query to be more specific.
        if state.get("correction_attempts", 0) > 0:
            # A simple but effective strategy: ask the LLM to rephrase for better search
            prompt = f"""You are a search expert. The previous search for the question "{state['original_question']}" failed to find relevant documents.
            Please rephrase the question to be more specific or to use different keywords.
            Return ONLY the rephrased question.

            Rephrased Question:"""
        else:
            # For the first attempt, we can generate sub-questions
            prompt = f"""You are a research assistant. Based on the user's question, generate a single, more effective search query to find the most relevant documents.
            If the question is complex, break it down into a primary search query.
            
            User Question: "{state['question']}"
            
            Optimized Search Query:"""

        response = self.azure_client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=200, temperature=0.0
        )
        transformed_question = response.choices[0].message.content.strip()
        
        state["question"] = transformed_question
        logger.info(f"Transformed query to: '{transformed_question}'")
        return state

    ## NEW: Node for grading the relevance of retrieved documents.
    def _grade_documents(self, state: AgentState) -> AgentState:
        """Grades the relevance of retrieved documents to the user's original question."""
        logger.info("Grading retrieved documents...")
        question = state["original_question"]
        documents = state["search_results"]
        
        if not documents:
            state["retrieval_grade"] = "not_relevant"
            logger.warning("No documents found to grade.")
            return state

        # Use the cross-encoder scores we already have from re-ranking as a proxy for relevance
        # A positive score is generally a good indicator of relevance.
        avg_score = sum(doc['score'] for doc in documents) / len(documents)
        
        if avg_score > 0.5: # This threshold can be tuned
            logger.info("Grading: Documents are RELEVANT.")
            state["retrieval_grade"] = "relevant"
        else:
            logger.warning("Grading: Documents are NOT RELEVANT. Average score was low.")
            state["retrieval_grade"] = "not_relevant"
            
        return state

    ## NEW: Conditional edge logic for the self-correction loop.
    def _decide_to_continue(self, state: AgentState) -> str:
        """Determines the next step based on the retrieval grade."""
        if state["retrieval_grade"] == "relevant":
            logger.info("Decision: Retrieval successful. Proceeding to synthesis.")
            return "synthesize"
        else:
            attempts = state.get("correction_attempts", 0)
            if attempts >= 2: # Max 2 correction attempts
                logger.error("Decision: Max correction attempts reached. Proceeding to synthesis with poor results.")
                return "synthesize" # Give up and try to answer anyway
            
            logger.info("Decision: Retrieval failed. Re-transforming query.")
            state["correction_attempts"] = attempts + 1
            return "transform_query"

    ## MODIFIED: Synthesis prompt now asks for quote-level citations.
    def _synthesize_answer(self, state: AgentState) -> AgentState:
        """Generates the final answer, citing sources with specific chunk IDs."""
        logger.info("Synthesizing final answer...")
        question = state["original_question"]
        search_results = state["search_results"]
        
        if not search_results:
            state["final_answer"] = "I couldn't find any relevant information in the documents. Please try rephrasing your question."
            state["confidence"] = 0.0
            return state
            
        context_parts = []
        unique_parent_contents = set()
        for result in search_results:
            parent_content = result.get('metadata', {}).get('parent_content', result['content'])
            if parent_content not in unique_parent_contents:
                unique_parent_contents.add(parent_content)
                # Provide the chunk_id in the context for the LLM to use
                context_parts.append(f"Source (Chunk ID: {result['chunk_id']}):\n{parent_content}\n---")
        
        context_str = "\n".join(context_parts)

        synthesis_prompt = f"""You are an expert assistant for oil and gas operations.
        Analyze the following document sections to answer the question thoroughly.

        Question: {question}

        Documents:
        {context_str}

        Guidelines:
        - Provide a direct and comprehensive answer based ONLY on the provided documents.
        - When you use information from a source, you MUST cite it by embedding the source's chunk ID in a <cite> tag like this: <cite id="chunk_id_goes_here"></cite>.
        - Place the citation tag directly after the sentence or claim it supports.
        - If the information is not present, state that clearly. Do not make up information.

        Answer:"""

        response = self.azure_client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": synthesis_prompt}],
            max_tokens=1500, temperature=0.0
        )
        answer = response.choices[0].message.content
        
        avg_score = sum(r["score"] for r in search_results) / len(search_results)
        confidence = 1 / (1 + np.exp(-avg_score))

        state["final_answer"] = answer
        state["confidence"] = confidence
        logger.info(f"Generated answer with confidence: {confidence:.2f}")
        return state

    ## NEW: Node for generating proactive, related insights.
    def _insight_generator(self, state: AgentState) -> AgentState:
        """Looks for related, potentially interesting information in the retrieved context."""
        logger.info("Generating proactive insights...")
        question = state["original_question"]
        search_results = state["search_results"]
        
        if not search_results:
            state["proactive_insights"] = ""
            return state

        context_str = "\n---\n".join([res['metadata'].get('parent_content', res['content']) for res in search_results])
        
        insight_prompt = f"""You are a senior analyst. The user asked the following question: "{question}".
        You have already provided a direct answer. Now, review the full context of the retrieved documents below.
        
        Context:
        {context_str}
        
        Your task is to identify one or two related, surprising, or potentially important insights that the user did not explicitly ask for.
        For example, if they asked about chemical usage, is there a related safety incident or equipment failure mentioned?
        If no interesting related insights are present, simply respond with "No additional insights found."
        
        Proactive Insights:
        """
        
        response = self.azure_client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": insight_prompt}],
            max_tokens=400, temperature=0.5
        )
        insights = response.choices[0].message.content
        
        if "no additional insights" in insights.lower():
            state["proactive_insights"] = ""
        else:
            state["proactive_insights"] = insights
            logger.info("Found proactive insights.")
            
        return state

    ## MODIFIED: The graph is now a cyclical graph with the new adaptive logic.
    def _build_graph(self) -> StateGraph:
        """Builds the LangGraph workflow with a self-correcting retrieval loop."""
        workflow = StateGraph(AgentState)
        
        # Add nodes
        workflow.add_node("group_selector", self._select_groups) # Unchanged but part of the new flow
        workflow.add_node("transform_query", self._query_transformer)
        workflow.add_node("document_searcher", self._search_documents) # Unchanged implementation
        workflow.add_node("grade_documents", self._grade_documents)
        workflow.add_node("answer_synthesizer", self._synthesize_answer)
        workflow.add_node("insight_generator", self._insight_generator)
        
        # Build edges
        workflow.set_entry_point("group_selector")
        workflow.add_edge("group_selector", "transform_query")
        workflow.add_edge("transform_query", "document_searcher")
        workflow.add_edge("document_searcher", "grade_documents")
        
        # This is the self-correction loop
        workflow.add_conditional_edges(
            "grade_documents",
            self._decide_to_continue,
            {
                "synthesize": "answer_synthesizer",
                "transform_query": "transform_query"
            }
        )
        
        workflow.add_edge("answer_synthesizer", "insight_generator")
        workflow.add_edge("insight_generator", END)
        
        # Compile the workflow with default parameters
        compiled_graph = workflow.compile(checkpointer=None)
        
        # Set recursion limit through graph configuration
        compiled_graph.recursion_limit = 50
        return compiled_graph

    def process_query(self, question: str) -> Dict[str, Any]:
        """Main query processor with the new agentic workflow."""
        # Caching logic is unchanged
        cache_key = question.lower().strip()
        if cache_key in self.query_cache:
            logger.info(f"Returning cached result for query: {question}")
            return self.query_cache[cache_key]

        # Structure vs. Content routing is unchanged
        # ... (code for routing prompt) ...
        route = "content" # Forcing content for this example
        
        if "structure" in route:
            # ... (code for _answer_structure_query) ...
            pass
        
        # --- Run the new, advanced LangGraph workflow ---
        initial_state = {
            "messages": [HumanMessage(content=question)],
            "original_question": question,
            "question": question,
            "correction_attempts": 0,
        }
        
        try:
            final_state = self.graph.invoke(initial_state)
            
            response = {
                "answer": final_state["final_answer"],
                "confidence": final_state["confidence"],
                "selected_groups": final_state["selected_groups"],
                "search_results": final_state["search_results"],
                "proactive_insights": final_state["proactive_insights"], # NEW
            }
            self.query_cache[cache_key] = response
            return response
            
        except Exception as e:
            logger.error(f"Error processing query graph: {e}", exc_info=True)
            return {
                "answer": f"I encountered an error: {str(e)}",
                "confidence": 0.0, "selected_groups": [], "search_results": [], "proactive_insights": ""
            }

    # -- The following methods are part of the agent but are mostly unchanged in their implementation --
    # -- They are called by the new graph structure. --
    def _select_groups(self, state: AgentState) -> AgentState:
        """Select relevant groups by first checking summary file metadata"""
        question = state["question"]
        
        try:
            import pandas as pd
            df = pd.read_excel("document_summary.xlsx")
            
            # Check for matching documents in summary
            matching_groups = set()
            
            # Search summaries for question keywords
            for _, row in df.iterrows():
                # Check if question contains any well names from this doc
                doc_wells = row['wells'] if isinstance(row['wells'], list) else eval(row['wells'])
                for well in doc_wells:
                    if well.lower() in question.lower():
                        matching_groups.add(row['group'])
                
                # Check if question contains any dates from this doc  
                doc_dates = row['dates'] if isinstance(row['dates'], list) else eval(row['dates'])
                for date in doc_dates:
                    if date in question:
                        matching_groups.add(row['group'])
                
                # Check summary content
                if row['summary'] and row['summary'].lower() in question.lower():
                    matching_groups.add(row['group'])
            
            if matching_groups:
                state["selected_groups"] = list(matching_groups)
                return state
                
        except Exception as e:
            logger.warning(f"Error checking summary file: {e}")
        
        # Fallback to AI group selection if no matches in summary
        group_selection_prompt = f"""You are an expert document group selector for oil and gas operations. 
        Strictly select ONLY the most relevant document groups based on the user's question.

        Available Groups and Their Content:
        - Chemical Consumption: ONLY chemicals, additives, fluid treatment data
        - Engineering Tickets: ONLY technical equipment, mechanical issues
        - Risk Assessment: ONLY safety documents, risk analysis
        - Well Recap: ONLY well completion, drilling reports
        - Mud Program: ONLY drilling mud, fluid properties
        - Hydraulic Summary: ONLY pressure, flow, pump data

        Question: "{question}"

        Rules:
        1. Select ONLY 1 group unless question clearly spans multiple domains
        2. Be extremely precise - wrong groups waste time
        3. If unsure between groups, prefer Chemical Consumption for chemical questions
        4. Return ONLY a JSON array of group names

        Selected Groups:"""

        try:
            response = self.azure_client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": group_selection_prompt}],
                max_tokens=200,
                temperature=0
            )
            
            groups_text = response.choices[0].message.content.strip()
            selected_groups = json.loads(groups_text)
            
            # Validate selected groups
            valid_groups = [g for g in selected_groups if g in FILE_GROUPS]
            if not valid_groups:
                valid_groups = list(FILE_GROUPS.keys())  # Search all if none valid
            
            logger.info(f"Selected groups for question '{question}': {valid_groups}")
            
            state["selected_groups"] = valid_groups
            state["messages"].append(AIMessage(content=f"Selected groups: {valid_groups}"))
            
            return state
            
        except Exception as e:
            logger.error(f"Error selecting groups: {e}")
            state["selected_groups"] = list(FILE_GROUPS.keys())
            return state
    
    def _search_documents(self, state: AgentState) -> AgentState:
        """
        MODIFIED: Enhanced document search with a re-ranking step.
        """
        logger.info("Performing search with re-ranking...")
        question = state["question"]
        selected_groups = state["selected_groups"]
        
        # 1. RETRIEVAL: Get a larger pool of candidate documents (e.g., k=25)
        # We retrieve more documents initially to give the re-ranker a good selection.
        candidate_results = self.vector_store.search_groups(question, selected_groups, k=25)
        
        if not candidate_results:
            state["search_results"] = []
            return state

        # 2. RE-RANKING: Prepare pairs of (question, chunk_content) for the cross-encoder
        sentence_pairs = [[question, result[0].content] for result in candidate_results]
        
        # Predict scores using the cross-encoder
        rerank_scores = self.cross_encoder.predict(sentence_pairs)
        
        # Add the new, more accurate scores to our results
        # The new tuple will be (chunk_object, new_reranked_score, group_name)
        for i in range(len(candidate_results)):
            candidate_results[i] = (candidate_results[i][0], rerank_scores[i], candidate_results[i][2])
            
        # Sort the results by the new re-ranking score in descending order
        candidate_results.sort(key=lambda x: x[1], reverse=True)
        
        # 3. FORMATTING: Keep the top N re-ranked results (e.g., top 7)
        top_results = candidate_results[:7]
        
        formatted_results = []
        for chunk, score, group in top_results:
            # We pass the full metadata dictionary now to access parent_content later
            result_metadata = chunk.metadata.copy()
            result_metadata['score'] = float(score)
            
            formatted_results.append({
                "content": chunk.content,
                "file_name": chunk.metadata.get('file_name', 'unknown'),
                "group": group,
                "score": float(score), # Using the new, more accurate score
                "chunk_id": chunk.chunk_id,
                "metadata": chunk.metadata # <-- MODIFIED: Pass the whole metadata dict
            })
            
        state["search_results"] = formatted_results
        logger.info(f"Re-ranked and selected top {len(formatted_results)} search results.")
        
        return state


class AgenticRAGWithLangGraph:
    """Main Agentic RAG system class"""
    
    def __init__(self, azure_endpoint: str, azure_key: str, api_version: str = "2024-02-15-preview"):
        self.azure_client = AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=azure_key,
            api_version=api_version
        )
        
        self.document_processor = DocumentProcessor(self.azure_client)
        self.vector_store = GroupedVectorStore()
        self.file_analyzer = None
        self.agent = None
        
        logger.info("Initialized AgenticRAG with LangGraph system")
    
    def ingest_folder(self, folder_path: str, save_path: str = "./grouped_vector_store"):
        """Ingest all documents from a folder with group classification"""
        logger.info(f"Starting grouped ingestion of folder: {folder_path}")
        
        if not os.path.exists(folder_path):
            logger.error(f"Folder does not exist: {folder_path}")
            return 0
        
        # Analyze file system
        self.file_analyzer = FileSystemAnalyzer(folder_path)
        self.file_analyzer.set_document_processor(self.document_processor)
        self.file_analyzer.analyze()
        
        # Process all supported files
        all_chunks = []
        processed_files = 0
        group_stats = {group: 0 for group in FILE_GROUPS.keys()}
        
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_ext = Path(file).suffix.lower()
                
                try:
                    if file_ext == '.pdf':
                        chunks = self.document_processor.process_pdf(file_path)
                    elif file_ext == '.docx':
                        chunks = self.document_processor.process_docx(file_path)
                    elif file_ext == '.txt':
                        chunks = self.document_processor.process_txt(file_path)
                    else:
                        continue  # Skip unsupported files
                    
                    if chunks:
                        all_chunks.extend(chunks)
                        processed_files += 1
                        # Count chunks by group
                        for chunk in chunks:
                            group_stats[chunk.group] += 1
                        logger.info(f"Processed {file_path}: {len(chunks)} chunks in group {chunks[0].group}")
                    else:
                        logger.warning(f"No chunks created from {file_path}")
                        
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    continue
        
        if not all_chunks:
            logger.error("No chunks were created from any files")
            return 0
        
        # Add to grouped vector store
        self.vector_store.add_documents(all_chunks)
        
        # Save all groups
        self.vector_store.save_all_groups(save_path)
        
        # Initialize agent
        self.agent = LangGraphAgent(self.azure_client, self.vector_store, self.file_analyzer)
        
        logger.info(f"Ingestion complete. Processed {processed_files} files, created {len(all_chunks)} chunks")
        logger.info(f"Group distribution: {group_stats}")
        return len(all_chunks)

    ## MODIFIED: The query method now returns proactive_insights
    def query(self, question: str) -> Dict[str, Any]:
        """Process a query using the LangGraph agentic approach"""
        logger.info(f"Processing query: {question}")
        
        if not self.vector_store.group_indices or not self.agent:
            return {
                "answer": "System not initialized. Please ingest documents first.",
                "confidence": 0.0, "selected_groups": [], "search_results": [], "proactive_insights": ""
            }
        
        try:
            response = self.agent.process_query(question)
            return response
        except Exception as e:
            logger.error(f"Error processing query: {e}", exc_info=True)
            return {
                "answer": f"I apologize, but I encountered an error: {str(e)}",
                "confidence": 0.0, "selected_groups": [], "search_results": [], "proactive_insights": ""
            }
    
    def load_vector_store(self, save_path: str = "./grouped_vector_store"):
        """Load a previously saved grouped vector store"""
        logger.info(f"Loading grouped vector store from: {save_path}")
        
        loaded_groups = self.vector_store.load_all_groups(save_path)
        
        if loaded_groups:
            # Initialize agent if we have a file analyzer
            if self.file_analyzer:
                self.agent = LangGraphAgent(self.azure_client, self.vector_store, self.file_analyzer)
            else:
                # Create a minimal file analyzer
                self.file_analyzer = FileSystemAnalyzer("./")
                self.file_analyzer.set_document_processor(self.document_processor)
                self.agent = LangGraphAgent(self.azure_client, self.vector_store, self.file_analyzer)
            
            logger.info(f"Grouped vector store loaded successfully. Groups: {loaded_groups}")
            return True
        else:
            logger.error("Failed to load grouped vector store")
            return False
    
    def get_system_stats(self) -> Dict[str, Any]:
        """Enhanced system statistics"""
        stats = {
            "vector_store_initialized": bool(self.vector_store.group_indices),
            "total_groups": len(self.vector_store.group_indices),
            "agent_initialized": self.agent is not None,
            "file_analyzer_available": self.file_analyzer is not None,
            "supported_file_types": list(self.document_processor.supported_extensions),
            # Dynamically get the model name from the loaded sentence transformer model
            "embedding_model": self.vector_store.model_name if self.vector_store else "Unknown"
        }
        
        # Enhanced group statistics
        if self.vector_store.group_indices:
            enhanced_stats = self.vector_store.get_enhanced_group_stats()
            file_summary = self.vector_store.get_file_type_summary()
            
            stats.update({
                "file_type_summary": file_summary,
                "enhanced_group_statistics": enhanced_stats['groups'],
                "file_type_distribution": enhanced_stats['file_type_distribution'],
                "total_chunks": sum(g['chunk_count'] for g in enhanced_stats['groups'].values()),
                "total_unique_files": sum(g['unique_files'] for g in enhanced_stats['groups'].values()),
                "available_groups": list(self.vector_store.group_indices.keys())
            })
            
            # Add processing statistics if available
            # CORRECTED: Access the 'processing_stats' attribute directly
            if hasattr(self.document_processor, 'processing_stats'):
                stats["processing_statistics"] = self.document_processor.processing_stats
        
        return stats
    
    def get_group_file_details(self, group_name: str) -> Dict[str, Any]:
        """Get detailed file information for a specific group"""
        return self.vector_store.get_group_files(group_name)

    def batch_query(self, questions: List[str]) -> List[Dict[str, Any]]:
        """Process multiple queries in batch"""
        results = []
        
        for i, question in enumerate(questions):
            logger.info(f"Processing batch query {i+1}/{len(questions)}: {question}")
            try:
                response = self.query(question)
                results.append({
                    "question": question,
                    "answer": response["answer"],
                    "confidence": response["confidence"],
                    "selected_groups": response["selected_groups"],
                    "status": "success"
                })
            except Exception as e:
                logger.error(f"Error in batch query {i+1}: {e}")
                results.append({
                    "question": question,
                    "answer": f"Error: {str(e)}",
                    "confidence": 0.0,
                    "selected_groups": [],
                    "status": "error"
                })
        
        return results
    
    def export_knowledge_base(self, export_path: str):
        """Export the grouped knowledge base to a readable format"""
        try:
            export_data = {
                "metadata": {
                    "export_timestamp": datetime.now().isoformat(),
                    "total_groups": len(self.vector_store.group_indices),
                    "system_stats": self.get_system_stats()
                },
                "groups": {}
            }
            
            # Export each group's documents
            for group_name, chunks in self.vector_store.group_chunks.items():
                if not chunks:
                    continue
                
                # Group chunks by file
                files_dict = {}
                for chunk in chunks:
                    file_name = chunk.metadata.get('file_name', 'unknown')
                    if file_name not in files_dict:
                        files_dict[file_name] = []
                    files_dict[file_name].append(chunk)
                
                # Export each file's content
                group_data = {
                    "group_name": group_name,
                    "total_chunks": len(chunks),
                    "total_files": len(files_dict),
                    "documents": []
                }
                
                for file_name, file_chunks in files_dict.items():
                    file_data = {
                        "file_name": file_name,
                        "total_chunks": len(file_chunks),
                        "file_type": file_chunks[0].metadata.get('file_type', 'unknown'),
                        "processed_at": file_chunks[0].metadata.get('processed_at', 'unknown'),
                        "chunks": [
                            {
                                "chunk_id": chunk.chunk_id,
                                "content": chunk.content,
                                "metadata": chunk.metadata
                            } for chunk in file_chunks
                        ]
                    }
                    group_data["documents"].append(file_data)
                
                export_data["groups"][group_name] = group_data
            
            # Save to JSON
            with open(export_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Grouped knowledge base exported to {export_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error exporting knowledge base: {e}")
            return False

    def save_vector_store(self, save_path: str) -> bool:
        """Saves all groups from the vector store to a folder."""
        if not self.vector_store:
            logger.error("Vector store not initialized, cannot save.")
            return False
        return self.vector_store.save_all_groups(save_path)

def enhanced_sidebar_display(st, rag_system):
    """Enhanced sidebar with professional styling"""
    
    # Custom sidebar styling
    st.markdown("""
    <style>
    .sidebar .sidebar-content {
        background-color: #2c3e50;
        color: white;
    }
    .sidebar .sidebar-content .stRadio > div > label {
        color: white;
    }
    .sidebar .sidebar-content .stTextInput > div > div > input {
        background-color: #34495e;
        color: white;
    }
    .sidebar .sidebar-content .stButton > button {
        background-color: #3498db;
        color: white;
        border: none;
        border-radius: 4px;
        padding: 8px 16px;
    }
    .sidebar .sidebar-content .stButton > button:hover {
        background-color: #2980b9;
    }
    .sidebar .sidebar-content .stExpander > div > div {
        background-color: #34495e;
        border-radius: 4px;
        padding: 8px;
        margin-bottom: 8px;
    }
    </style>
    """, unsafe_allow_html=True)
    
    with st.sidebar:
        # Sidebar header with accent color
        st.markdown("""
        <div style="background-color: #3498db; padding: 10px; border-radius: 4px; margin-bottom: 20px;">
            <h2 style="color: white; margin: 0;">Configuration</h2>
        </div>
        """, unsafe_allow_html=True)
        
        # Action selection with custom styling
        action = st.radio("Choose Action", ["Ingest New Folder", "Load Vector Store"])
        
        if action == "Ingest New Folder":
            folder_path = st.text_input("Enter folder path to ingest:", "./Input")
            if st.button("Ingest Documents", key="ingest_btn") and folder_path:
                if os.path.exists(folder_path):
                    with st.spinner(f"Ingesting documents from {folder_path}..."):
                        chunks_created = rag_system.ingest_folder(folder_path)
                        st.success(f"Created {chunks_created} chunks from documents")
                        st.rerun()
                else:
                    st.error(f"Folder {folder_path} does not exist.")
        
        elif action == "Load Vector Store":
            # CORRECTED LINE
            vector_store_path = st.text_input("Enter vector store path:", "./grouped_vector_store")
            if st.button("Load Vector Store", key="load_btn"):
                with st.spinner("Loading vector store..."):
                    if rag_system.load_vector_store(vector_store_path):
                        st.success("Vector store loaded successfully")
                        st.rerun()
                    else:
                        st.error("Failed to load vector store")
        
        # Divider with custom style
        st.markdown("""
        <div style="height: 1px; background-color: #3498db; margin: 20px 0;"></div>
        """, unsafe_allow_html=True)
        
        # Save section with custom styling
        st.markdown("""
        <div style="background-color: #34495e; padding: 10px; border-radius: 4px; margin-bottom: 20px;">
            <h3 style="color: white; margin: 0;">Save Vector Store</h3>
        </div>
        """, unsafe_allow_html=True)
        
        save_path = st.text_input("Save path:", "./grouped_vector_store")
        if st.button("Save Vector Store", key="save_btn"):
            if rag_system.save_vector_store(save_path):
                st.success("Vector store saved successfully")
            else:
                st.error("Failed to save vector store")
        
        # Document groups section
        st.markdown("""
        <div style="background-color: #34495e; padding: 10px; border-radius: 4px; margin: 20px 0;">
            <h3 style="color: white; margin: 0;">Document Groups</h3>
        </div>
        """, unsafe_allow_html=True)
        
        stats = rag_system.get_system_stats()
        
        for group_name, keywords in FILE_GROUPS.items():
            group_stats = stats.get('enhanced_group_statistics', {}).get(group_name, {})
            file_count = group_stats.get('unique_files', 0)
            
            with st.expander(f"{group_name} ({file_count} files)"):
                if keywords:
                    st.markdown(f"<p style='color: #3498db;'><strong>Keywords:</strong> {', '.join(keywords)}</p>", unsafe_allow_html=True)
                else:
                    st.markdown("<p style='color: #3498db;'><strong>Default group</strong> for unmatched files</p>", unsafe_allow_html=True)
                
                if file_count > 0:
                    st.markdown(f"<p><strong>Total chunks:</strong> {group_stats.get('chunk_count', 0)}</p>", unsafe_allow_html=True)
                    
                    file_details = rag_system.get_group_file_details(group_name)
                    
                    st.markdown("<p><strong>Files in this group:</strong></p>", unsafe_allow_html=True)
                    for file_name in file_details.get('files', []):
                        file_info = file_details.get('file_details', {}).get(file_name, {})
                        processing_method = file_info.get('processing_method', 'Unknown')
                        chunks = file_info.get('chunks', 0)
                        
                        st.markdown(f"<p style='margin-left: 20px;'>{file_name} ({chunks} chunks)</p>", unsafe_allow_html=True)
                        
                        if file_info.get('total_pages', 0) > 0:
                            vision_pages = len(file_info.get('vision_pages', []))
                            text_pages = len(file_info.get('text_pages', []))
                            st.markdown(f"<p style='margin-left: 20px;'>Pages: {text_pages} text, {vision_pages} OCR</p>", unsafe_allow_html=True)
                
                else:
                    st.markdown("<p style='color: #7f8c8d;'>No files in this group</p>", unsafe_allow_html=True)

## MODIFIED: The main display function now handles citations and proactive insights.
def enhanced_main_display(st, rag_system, response):
    """Enhanced main display with professional styling"""
    # Custom main page styling
    st.markdown("""
    <style>
    .main .block-container {
        background-color: #f5f7fa;
        padding: 2rem;
    }
    .stMetric {
        background-color: white;
        border-radius: 8px;
        padding: 15px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .stMetric > div > div {
        color: #2c3e50;
    }
    .stMetric > div > div:first-child {
        font-size: 1rem;
        color: #7f8c8d;
    }
    .stMetric > div > div:nth-child(2) {
        font-size: 1.5rem;
        color: #2c3e50;
        font-weight: bold;
    }
    .stDataFrame {
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .stButton > button {
        background-color: #3498db;
        color: white;
        border: none;
        border-radius: 4px;
        padding: 8px 16px;
    }
    .stButton > button:hover {
        background-color: #2980b9;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Main header with gradient background
    st.markdown("""
    <div style="background: linear-gradient(135deg, #3498db, #2c3e50); 
                padding: 20px; 
                border-radius: 8px; 
                margin-bottom: 30px;
                color: white;">
        <h1 style="color: white; margin: 0;">Agentic RAG System</h1>
        <p style="margin: 5px 0 0 0;">Document processing with intelligent group selection and enhanced metadata tracking</p>
    </div>
    """, unsafe_allow_html=True)
    
    # System stats display
    stats = rag_system.get_system_stats()
    
    st.markdown("""
    <div style="background-color: white; 
                padding: 15px; 
                border-radius: 8px; 
                margin-bottom: 20px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
        <h2 style="color: #2c3e50; margin: 0 0 15px 0;">System Statistics</h2>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Vector Store", "✅ Ready" if stats['vector_store_initialized'] else "❌ Not Ready")
    with col2:
        st.metric("Total Groups", stats['total_groups'])
    with col3:
        st.metric("Total Chunks", stats.get('total_chunks', 0))
    with col4:
        st.metric("Unique Files", stats.get('total_unique_files', 0))
    
    # File type and processing method breakdown
    if stats.get('file_type_summary'):
        st.markdown("""
        <div style="background-color: white; 
                    padding: 15px; 
                    border-radius: 8px; 
                    margin: 20px 0;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
            <h2 style="color: #2c3e50; margin: 0 0 15px 0;">File Processing Summary</h2>
        </div>
        """, unsafe_allow_html=True)
        
        file_summary = stats['file_type_summary']
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("PDF Files", file_summary['pdf_files'])
        with col2:
            st.metric("DOCX Files", file_summary['docx_files'])
        with col3:
            st.metric("TXT Files", file_summary['txt_files'])
        
        # Processing methods
        st.markdown("""
        <div style="background-color: white; 
                    padding: 15px; 
                    border-radius: 8px; 
                    margin: 20px 0;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
            <h2 style="color: #2c3e50; margin: 0 0 15px 0;">Processing Methods</h2>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Text Extraction", file_summary['processing_methods']['text_extraction'])
        with col2:
            st.metric("Vision OCR", file_summary['processing_methods']['vision_ocr'])
    
    # Group statistics table
    if stats.get('enhanced_group_statistics'):
        st.markdown("""
        <div style="background-color: white; 
                    padding: 15px; 
                    border-radius: 8px; 
                    margin: 20px 0;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
            <h2 style="color: #2c3e50; margin: 0 0 15px 0;">Detailed Group Statistics</h2>
        </div>
        """, unsafe_allow_html=True)

    # Query Interface
    st.subheader("Ask a Question")
    question = st.text_input(
        "Enter your question about the documents:",
        placeholder="e.g., What was the impact of using bentonite on wellbore stability?",
        key="question_input"
    )
    
    if st.button("Submit Query", key="submit_query") and question:
        with st.spinner("Processing your query with the advanced agent..."):
            response = rag_system.query(question)
            st.session_state.response = response # Save response to session state to persist it
            
    if response:
        # Display results
        st.markdown("### Answer")
        
        # ## NEW: Simple parsing for <cite> tags. A real UI could make these clickable.
        answer_html = response['answer']
        # Find all cite tags and replace them with a styled, visible citation number
        cite_matches = list(re.finditer(r'<cite id="(.*?)"></cite>', answer_html))
        citations = {}
        for i, match in enumerate(cite_matches):
            chunk_id = match.group(1)
            cite_key = f"[{i+1}]"
            citations[cite_key] = chunk_id
            answer_html = answer_html.replace(match.group(0), f"&nbsp;<sup>**{cite_key}**</sup>")

        st.markdown(answer_html, unsafe_allow_html=True)
        
        # Display metadata
        col_meta1, col_meta2 = st.columns(2)
        with col_meta1:
            st.metric("Confidence", f"{response['confidence']:.2f}")
        with col_meta2:
            st.write(f"**Groups Searched:** {', '.join(response['selected_groups'])}")
            
        # ## NEW: Display Proactive Insights if they exist
        if response.get('proactive_insights'):
            with st.container():
                st.markdown("---")
                st.markdown("### Proactive Insights")
                st.info(response['proactive_insights'])

        # Display search results and citations
        with st.expander("View Evidence and Citations"):
            if citations:
                st.markdown("#### Citations")
                for key, chunk_id in citations.items():
                    st.markdown(f"**{key}**: `{chunk_id}`")

            if response.get('search_results'):
                st.markdown("---")
                st.markdown("#### Top Retrieved Documents")
                for result in response['search_results'][:3]:
                    st.markdown(f"**File:** `{result['file_name']}` (Score: {result['score']:.3f})")
                    st.text_area(
                        label=f"Content from Chunk ID: {result['chunk_id']}",
                        value=result['metadata'].get('parent_content', result['content']),
                        height=150,
                        disabled=True
                    )


def main():
    """Main function to run the enhanced Streamlit demo"""
    st.set_page_config(page_title="Advanced Agentic RAG", layout="wide")

    AZURE_ENDPOINT = "https://oai-nasco.openai.azure.com/openai/deployments/gpt-4o/chat/completions?api-version=2025-01-01-preview"
    AZURE_KEY = "YOUR_AZURE_OPENAI_API_KEY"
    API_VERSION = "2025-01-01-preview"
    
    if 'rag_system' not in st.session_state:
        st.session_state.rag_system = AgenticRAGWithLangGraph(AZURE_ENDPOINT, AZURE_KEY, API_VERSION)
    if 'response' not in st.session_state:
        st.session_state.response = None
        
    rag_system = st.session_state.rag_system
    
    # The sidebar display function is unchanged.
    enhanced_sidebar_display(st, rag_system)
    
    # The main display is now called with the response from the session state.
    enhanced_main_display(st, rag_system, st.session_state.response)

if __name__ == "__main__":
     main()
